from collections import defaultdict
from datetime import date, datetime
from decimal import Decimal
from io import BytesIO
import json
from pathlib import Path

from django.conf import settings
from django.db.models import Count
from django.utils import timezone
from docx import Document as WordDocument
from docx.enum.section import WD_ORIENTATION
from docx.shared import Inches, Pt
from openpyxl import Workbook
from openpyxl.chart import BarChart, LineChart, PieChart, Reference
try:
    from openpyxl.drawing.image import Image as OpenPyxlImage
except ImportError:  # pragma: no cover - optional image support depends on Pillow.
    OpenPyxlImage = None
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
from openpyxl.utils import get_column_letter
from reportlab.lib import colors
from reportlab.lib.pagesizes import landscape, letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

from apps.common.models import DuplicateReviewQueue
from apps.dashboard.models import Receipt
from apps.workforce.models import (
    Application,
    CommunityHealthWorker,
    DataImportBatch,
    HealthStudent,
    ImportedWorkbookSheet,
    MedicalDoctor,
    MissingDataReview,
    Midwife,
    NurseAide,
    NursingProfessional,
    PracticingLicenseRecord,
)


NURSING_TARGETS = ["nursingprofessional", "midwife", "nurseaide", "healthstudent"]
MEDICAL_TARGETS = ["medicaldoctor", "communityhealthworker"]
NURSING_FORMS = ["G1", "G2", "G3", "G4", "G5", "G6", "G7", "NC1", "NC2", "NC3", "NC4", "NC5", "NC6", "NC7", "NC8", "NC9", "NC10", "NC11"]
MEDICAL_FORMS = ["MD1", "MD2", "CHW1", "MBSP", "MBRN", "MBAC", "MBPF", "MBTC"]
GOV_NAVY = "12324A"
GOV_GREEN = "0F766E"
GOV_GOLD = "B0892F"
GOV_LIGHT = "F4F7F6"
GOV_BORDER = "C7D2D0"
PDF_AVAILABLE_WIDTH = landscape(letter)[0] - 48
FINANCIAL_MIN_VALID_DATE = date(2000, 1, 1)


def _ndoh_logo_path():
    logo = settings.BASE_DIR / "static" / "img" / "NDOH_LOGO.png"
    return logo if logo.exists() else None


def _month_key(record):
    source_date = record.issued_date or record.payment_date
    if not source_date:
        source_date = timezone.localtime(record.created_at).date() if timezone.is_aware(record.created_at) else record.created_at.date()
    return source_date.strftime("%Y-%m")


def _office_config(office):
    if office == "medical":
        return {
            "label": "Medical Board",
            "targets": MEDICAL_TARGETS,
            "forms": MEDICAL_FORMS,
            "batch_kinds": ["medical_board_workbook"],
        }
    return {
        "label": "Nursing Council",
        "targets": NURSING_TARGETS,
        "forms": NURSING_FORMS,
        "batch_kinds": ["ndata_workbook", "nursing_full_registration_2026"],
    }


def _normalize_office_selection(office=None):
    if office is None:
        return ["nursing", "medical"]
    if isinstance(office, (list, tuple, set)):
        offices = [item for item in office if item in {"nursing", "medical"}]
    else:
        offices = [office] if office in {"nursing", "medical"} else []
    return offices or ["nursing", "medical"]


def _count_by(rows, key_name, default="Not captured", limit=15):
    counts = defaultdict(int)
    for row in rows:
        label = getattr(row, key_name, "") or default
        counts[label] += 1
    return sorted(counts.items(), key=lambda item: item[1], reverse=True)[:limit]


def _registry_models_for_office(office_key):
    if office_key == "medical":
        return [
            ("Medical Doctors", MedicalDoctor),
            ("Community Health Workers", CommunityHealthWorker),
        ]
    return [
        ("Registered Nurses", NursingProfessional),
        ("Midwives", Midwife),
        ("Nurse Aides", NurseAide),
        ("Graduands", HealthStudent),
    ]


def _format_timestamp(value):
    if not value:
        return ""
    if timezone.is_aware(value):
        value = timezone.localtime(value)
    return value.strftime("%Y-%m-%d %H:%M")


def _build_guide_rows():
    return [
        [
            "Read this workbook in order",
            "Start with Executive Summary, then How To Read, then the office Overview sheet, and only after that move into Monthly, Yearly, Mix, and Admin tabs.",
            "This order helps leadership understand the source, the meaning, and the corrected totals before looking at detailed tables.",
        ],
        [
            "Current Registry Totals",
            "These are current live person counts from the main registry tables such as Midwife, NursingProfessional, MedicalDoctor, and CHW.",
            "Use this section when management asks how many people are currently in the live registry.",
        ],
        [
            "Imported Record Rows",
            "These are row-level operational records from PracticingLicenseRecord, including full registration, practising licence, provisional, temporary, workforce listing, and payment rows.",
            "This is not the same as current people because one person may contribute more than one imported row over time.",
        ],
        [
            "Monthly Analytics",
            "Each month is grouped from issued_date, payment_date, or created_at when no issue/payment date exists.",
            "Use this to explain the monthly trend of records processed or imported.",
        ],
        [
            "Unique People",
            "This is a distinct count using registration number, practitioner number, or full name within each month.",
            "This helps separate people counts from raw transaction row counts.",
        ],
        [
            "Applications By Status",
            "These counts come from the Application table and are filtered by office-specific form codes.",
            "Use this to explain how many submissions are pending, approved, or rejected.",
        ],
        [
            "Recent Import Batches",
            "These identify the last source file, source kind, completion date, and row volumes loaded into the system.",
            "Use this when management asks where the data came from and how recent it is.",
        ],
        [
            "Live Update vs Source Date",
            "A live registry count can be updated after import because staff may clean or edit records in the main tables.",
            "Explain the source month separately from the latest live update timestamp so management sees both dates clearly.",
        ],
        [
            "Example: Midwives",
            "The dashboard figure 2015 comes from Midwife.objects.count() in the live registry, not from a single April worksheet total.",
            "Use the office overview sheet to show the source file, latest import date, and the live table calculation behind the total.",
        ],
        [
            "Corrected Historical Totals",
            "If older exports were generated before duplicate cleansing or before office scoping was corrected, those totals may appear higher than the current report.",
            "Treat the current report as the corrected version and use the cleansing notes to explain why older totals changed.",
        ],
        [
            "Current Registry People vs Imported Record Rows",
            "Current Registry People are live people records. Imported Record Rows are historical activity rows such as renewals, provisional entries, workforce listings, and other imported transactions.",
            "Do not add these two together. They answer different questions.",
        ],
        [
            "Application Status Totals",
            "Approved, Pending, and Rejected come from the live Application table, filtered only to the selected office scope.",
            "Use these to explain the current workflow backlog and approval status, not historical import volume.",
        ],
    ]


def _build_workbook_navigation_rows():
    return [
        ["1", "Executive Summary", "High-level headline figures for the selected office.", "Use this page first for briefing notes and management summaries."],
        ["2", "How To Read", "Definitions and interpretation guidance for each type of statistic.", "Use this when explaining the figures to management or the Minister."],
        ["3", "Overview", "Live registry counts, corrected imported totals, source file details, and cleansing notes.", "Use this sheet to explain where the numbers came from and why they are valid."],
        ["4", "Monthly", "Detailed month-by-month history of imported activity rows.", "Use this to explain monthly trends and movement."],
        ["5", "Yearly", "Year-by-year imported activity totals.", "Use this for longer historical trend reporting."],
        ["6", "Mix", "Record-type and category mix tables/charts.", "Use this to explain what type of activity makes up the total."],
        ["7", "Admin", "Applications by status and recent import batch details.", "Use this to explain workflow status and the latest source loads."],
    ]


def _source_sheet_summary(rows, limit=4):
    if not rows:
        return "No workbook sheet details captured"
    return ", ".join(
        f"{row['sheet_name'].strip()} ({row['imported_rows']} rows)"
        for row in rows[:limit]
    )


def _build_comparison_rows(data):
    latest_batch = data.get("latest_batch_row") or {}
    latest_source_file = latest_batch.get("source_file_name") or "Not captured"
    latest_completed = latest_batch.get("completed_at") or "Not captured"
    recent_sheets = _source_sheet_summary(data.get("source_sheet_rows", []))
    difference = data["total_records"] - data["current_registry_total"]
    return [
        [
            "Current Registry People",
            data["current_registry_total"],
            "Current live person records stored in the main registry tables for this office.",
        ],
        [
            "Imported Record Rows",
            data["total_records"],
            "Operational/import rows from PracticingLicenseRecord. One person can contribute more than one row over time.",
        ],
        [
            "Difference Between Rows and People",
            difference,
            "This difference is expected because renewals, payments, workforce entries, and historical records increase row counts.",
        ],
        [
            "Latest Completed Import",
            latest_completed,
            f"Most recent completed source load for this office from {latest_source_file}.",
        ],
        [
            "Recent Source Sheets",
            recent_sheets,
            "Recent workbook tabs that fed the latest completed import batch.",
        ],
    ]


def _build_registry_explanation_rows(data):
    latest_batch = data.get("latest_batch_row") or {}
    latest_source_file = latest_batch.get("source_file_name") or "Not captured"
    latest_completed = latest_batch.get("completed_at") or "Not captured"
    recent_sheets = _source_sheet_summary(data.get("source_sheet_rows", []))
    rows = []
    for row in data["current_registry_rows"]:
        rows.append([
            row["label"],
            row["current_count"],
            row["active_count"],
            row["source_model"],
            f"{row['source_model']}.objects.count()",
            row["latest_updated"],
            latest_completed,
            latest_source_file,
            recent_sheets,
        ])
    return rows


def _build_management_answer_rows(data):
    latest_batch = data.get("latest_batch_row") or {}
    recent_sheets = _source_sheet_summary(data.get("source_sheet_rows", []))
    rows = [
        [
            "Where did this office's report data come from?",
            "From the live registry tables, the imported workbook batches, the application table, and the practising licence record table.",
            f"Latest completed source file: {latest_batch.get('source_file_name') or 'Not captured'} on {latest_batch.get('completed_at') or 'Not captured'}.",
        ],
        [
            "What is the most recent source period shown here?",
            recent_sheets,
            "Look for the latest month-named workbook tab in the recent source sheets list.",
        ],
    ]
    for row in data["current_registry_rows"]:
        rows.append([
            f"How was the current {row['label']} total calculated?",
            f"{row['current_count']} current records in the live registry.",
            f"Counted directly from the {row['source_model']} table. Latest live update recorded at {row['latest_updated'] or 'Not captured'}.",
        ])
    return rows


def _latest_duplicate_cleanup_summary():
    reports_dir = Path(__file__).resolve().parents[2] / "docs" / "reports"
    if not reports_dir.exists():
        return None
    candidates = sorted(reports_dir.glob("duplicate_cleanup_report_*.json"), reverse=True)
    for path in candidates:
        try:
            return json.loads(path.read_text(encoding="utf-8"))
        except (OSError, json.JSONDecodeError):
            continue
    return None


def _record_type_count_map(data):
    return {row["record_type"]: row["count"] for row in data.get("record_type_rows", [])}


def _application_status_count_map(data):
    return {row["status"]: row["count"] for row in data.get("application_rows", [])}


def _build_cleansing_rows(data, cleanup_summary=None):
    rows = [
        [
            "Current Nursing Council analytics scope",
            "Only Nursing Council target models are counted here: registered nurses, midwives, nurse aides, and graduands.",
            "This excludes Medical Board and CHW records from the Nursing Council monthly analytics totals.",
        ],
    ]
    if cleanup_summary:
        rows.extend([
            [
                "Exact duplicate groups identified in imported history",
                cleanup_summary.get("exact_duplicate_group_count", ""),
                "These were exact repeated imported rows found in the spreadsheet history table before cleanup.",
            ],
            [
                "Exact duplicate imported rows deleted",
                cleanup_summary.get("exact_duplicate_rows_deleted", ""),
                "Only imported historical duplicates were removed. No live practitioner registry rows were deleted.",
            ],
            [
                "Imported history rows remaining after cleanup",
                cleanup_summary.get("practicing_license_record_total", ""),
                "This is the current size of the imported historical backend after duplicate cleansing.",
            ],
            [
                "Pending duplicate review queue after cleanup",
                cleanup_summary.get("pending_review_queue_after", ""),
                "These are remaining suspicious groups for registrar/admin review rather than automatic deletion.",
            ],
        ])
    rows.append(
        [
            "Why some older reports may look higher",
            "Older generated reports may have been created before duplicate cleansing or before the Nursing Council scope was tightened.",
            "Use the current corrected record activity totals below when preparing management or ministerial submissions.",
        ]
    )
    return rows


def _build_nursing_correction_rows(data):
    record_counts = _record_type_count_map(data)
    application_counts = _application_status_count_map(data)
    return [
        [
            "Approved Nursing Council applications",
            application_counts.get("approved", 0),
            "Counted from the live Application table using Nursing Council form codes only.",
        ],
        [
            "Pending Nursing Council applications",
            application_counts.get("pending", 0),
            "Counted from the live Application table using Nursing Council form codes only.",
        ],
        [
            "Rejected Nursing Council applications",
            application_counts.get("rejected", 0),
            "Counted from the live Application table using Nursing Council form codes only.",
        ],
        [
            "Full registration activity rows",
            record_counts.get("full", 0),
            "Counted from PracticingLicenseRecord for Nursing Council target models after duplicate cleansing.",
        ],
        [
            "Practising licence activity rows",
            record_counts.get("practicing_license", 0),
            "Counted from PracticingLicenseRecord for Nursing Council target models after duplicate cleansing.",
        ],
        [
            "Provisional activity rows",
            record_counts.get("provisional", 0),
            "Counted from PracticingLicenseRecord for Nursing Council target models after duplicate cleansing.",
        ],
        [
            "Temporary activity rows",
            record_counts.get("temporary", 0),
            "Counted from PracticingLicenseRecord for Nursing Council target models after duplicate cleansing.",
        ],
        [
            "Workforce listing activity rows",
            record_counts.get("workforce_listing", 0),
            "Counted from PracticingLicenseRecord for Nursing Council target models after duplicate cleansing.",
        ],
        [
            "Current corrected Nursing Council imported activity total",
            data.get("total_records", 0),
            "This is the corrected current Nursing Council analytics total used for monthly reporting.",
        ],
    ]


def build_monthly_analytics_payload(office=None):
    office_keys = _normalize_office_selection(office)
    offices = {}
    for office_key in office_keys:
        config = _office_config(office_key)
        records = list(
            PracticingLicenseRecord.objects.filter(target_model__in=config["targets"])
            .order_by("record_year", "issued_date", "created_at")
        )
        monthly = defaultdict(lambda: {
            "month": "",
            "full": 0,
            "practicing_license": 0,
            "provisional": 0,
            "temporary": 0,
            "workforce_listing": 0,
            "payment": 0,
            "total_records": 0,
            "unique_people": set(),
        })
        yearly = defaultdict(lambda: {
            "year": "",
            "full": 0,
            "practicing_license": 0,
            "provisional": 0,
            "temporary": 0,
            "workforce_listing": 0,
            "payment": 0,
            "total_records": 0,
        })

        for record in records:
            month = _month_key(record)
            year = record.record_year or int(month[:4])
            person_key = record.registration_no or record.practitioner_number or record.full_name
            monthly[month]["month"] = month
            monthly[month][record.record_type] = monthly[month].get(record.record_type, 0) + 1
            monthly[month]["total_records"] += 1
            if person_key:
                monthly[month]["unique_people"].add(person_key)
            yearly[year]["year"] = year
            yearly[year][record.record_type] = yearly[year].get(record.record_type, 0) + 1
            yearly[year]["total_records"] += 1

        monthly_rows = []
        for row in sorted(monthly.values(), key=lambda item: item["month"]):
            row = dict(row)
            row["unique_people"] = len(row["unique_people"])
            monthly_rows.append(row)

        yearly_rows = [dict(row) for row in sorted(yearly.values(), key=lambda item: item["year"])]
        record_type_rows = list(
            PracticingLicenseRecord.objects.filter(target_model__in=config["targets"])
            .values("record_type")
            .annotate(count=Count("id"))
            .order_by("record_type")
        )
        application_rows = list(
            Application.objects.filter(form_code__in=config["forms"])
            .values("status")
            .annotate(count=Count("id"))
            .order_by("status")
        )
        batch_queryset = DataImportBatch.objects.filter(source_kind__in=config["batch_kinds"])
        batch_rows = list(
            batch_queryset
            .values("id", "source_file_name", "source_kind", "status", "processed_rows", "total_rows", "completed_at")
            .order_by("-started_at")[:10]
        )
        latest_batch = batch_queryset.filter(status="completed").order_by("-completed_at", "-started_at").first()
        source_sheet_rows = []
        if latest_batch:
            source_sheet_rows = list(
                ImportedWorkbookSheet.objects.filter(batch=latest_batch)
                .values("sheet_name", "sheet_type", "status", "imported_rows", "skipped_rows")
                .order_by("id")[:20]
            )

        current_registry_rows = []
        current_registry_total = 0
        for label, model in _registry_models_for_office(office_key):
            current_count = model.objects.count()
            active_count = model.objects.filter(is_active=True).count()
            latest_updated = model.objects.order_by("-updated_at").values_list("updated_at", flat=True).first()
            current_registry_total += current_count
            current_registry_rows.append({
                "label": label,
                "current_count": current_count,
                "active_count": active_count,
                "latest_updated": _format_timestamp(latest_updated),
                "source_model": model.__name__,
            })

        offices[office_key] = {
            "label": config["label"],
            "total_records": len(records),
            "current_registry_total": current_registry_total,
            "current_registry_rows": current_registry_rows,
            "monthly_rows": monthly_rows,
            "yearly_rows": yearly_rows,
            "record_type_rows": record_type_rows,
            "category_rows": _count_by(records, "category"),
            "province_rows": _count_by(records, "province"),
            "application_rows": application_rows,
            "batch_rows": batch_rows,
            "latest_batch_row": {
                "id": latest_batch.id,
                "source_file_name": latest_batch.source_file_name,
                "source_kind": latest_batch.source_kind,
                "status": latest_batch.status,
                "processed_rows": latest_batch.processed_rows,
                "total_rows": latest_batch.total_rows,
                "completed_at": _format_timestamp(latest_batch.completed_at),
                "started_at": _format_timestamp(latest_batch.started_at),
                "summary": latest_batch.summary,
            } if latest_batch else None,
            "source_sheet_rows": source_sheet_rows,
        }

    offices["combined"] = {"generated_on": date.today().strftime("%d %b %Y")}
    if len(office_keys) > 1:
        offices["combined"].update({
            "duplicate_reviews": DuplicateReviewQueue.objects.count(),
            "open_missing_reviews": MissingDataReview.objects.exclude(status="resolved").count(),
        })
    return offices


def _append_table(ws, title, headers, rows, *, freeze=True, auto_filter=True, spacer_rows=1):
    ws.append([title])
    title_row = ws.max_row
    ws.merge_cells(start_row=title_row, start_column=1, end_row=title_row, end_column=max(1, len(headers)))
    ws.cell(title_row, 1).font = Font(bold=True, size=14, color=GOV_NAVY)
    ws.cell(title_row, 1).fill = PatternFill("solid", fgColor=GOV_LIGHT)
    ws.cell(title_row, 1).alignment = Alignment(vertical="center")
    ws.row_dimensions[title_row].height = 24
    ws.append(headers)
    header_row = ws.max_row
    thin_border = Border(
        left=Side(style="thin", color=GOV_BORDER),
        right=Side(style="thin", color=GOV_BORDER),
        top=Side(style="thin", color=GOV_BORDER),
        bottom=Side(style="thin", color=GOV_BORDER),
    )
    for cell in ws[header_row]:
        cell.font = Font(bold=True, color="FFFFFF")
        cell.fill = PatternFill("solid", fgColor=GOV_GREEN)
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
        cell.border = thin_border
    for row in rows:
        ws.append(row)
    for row in ws.iter_rows(min_row=header_row + 1, max_row=ws.max_row, max_col=len(headers)):
        for cell in row:
            cell.alignment = Alignment(vertical="top", wrap_text=True)
            cell.border = thin_border
            if isinstance(cell.value, Decimal):
                cell.number_format = '#,##0.00'
            elif isinstance(cell.value, int):
                cell.number_format = '#,##0'
            text_length = len(str(cell.value or ""))
            if text_length > 100:
                ws.row_dimensions[cell.row].height = max(ws.row_dimensions[cell.row].height or 15, 48)
            elif text_length > 55:
                ws.row_dimensions[cell.row].height = max(ws.row_dimensions[cell.row].height or 15, 32)
    for column_cells in ws.columns:
        max_length = max(len(str(cell.value or "")) for cell in column_cells)
        ws.column_dimensions[get_column_letter(column_cells[0].column)].width = min(max(max_length + 2, 12), 36)
    if freeze and not ws.freeze_panes:
        ws.freeze_panes = ws.cell(header_row + 1, 1).coordinate
    if auto_filter:
        ws.auto_filter.ref = f"A{header_row}:{get_column_letter(len(headers))}{ws.max_row}"
    end_row = ws.max_row
    for _ in range(spacer_rows):
        ws.append([])
    return header_row, end_row


def _style_sheet(ws, tab_color=GOV_GREEN):
    ws.sheet_properties.tabColor = tab_color
    ws.page_setup.orientation = "landscape"
    ws.page_setup.fitToWidth = 1
    ws.page_setup.fitToHeight = 0
    ws.sheet_properties.pageSetUpPr.fitToPage = True
    ws.page_margins.left = 0.25
    ws.page_margins.right = 0.25
    ws.page_margins.top = 0.45
    ws.page_margins.bottom = 0.45
    ws.oddHeader.center.text = "&BNational Department of Health - Workforce Registry&B"
    ws.oddFooter.center.text = "Confidential Registry Analytics | Page &P of &N"


def _add_workbook_banner(ws, subtitle):
    ws.insert_rows(1, 4)
    logo_path = _ndoh_logo_path()
    if logo_path and OpenPyxlImage:
        try:
            logo = OpenPyxlImage(str(logo_path))
            logo.width = 52
            logo.height = 52
            ws.add_image(logo, "A1")
        except Exception:
            pass
    ws.merge_cells("A1:I1")
    ws["A1"] = "NATIONAL DEPARTMENT OF HEALTH"
    ws["A1"].font = Font(bold=True, size=16, color="FFFFFF")
    ws["A1"].fill = PatternFill("solid", fgColor=GOV_NAVY)
    ws["A1"].alignment = Alignment(horizontal="center", vertical="center")
    ws.merge_cells("A2:I2")
    ws["A2"] = "Workforce Online Registration System"
    ws["A2"].font = Font(bold=True, size=12, color=GOV_NAVY)
    ws["A2"].alignment = Alignment(horizontal="center")
    ws.merge_cells("A3:I3")
    ws["A3"] = subtitle
    ws["A3"].font = Font(italic=True, size=10, color=GOV_GOLD)
    ws["A3"].alignment = Alignment(horizontal="center")


def _add_line_chart(ws, title, min_row, max_row, data_cols, anchor):
    if max_row <= min_row:
        return
    chart = LineChart()
    chart.title = title
    chart.y_axis.title = "Count"
    chart.x_axis.title = "Month"
    for col in data_cols:
        data = Reference(ws, min_col=col, min_row=min_row, max_row=max_row)
        chart.add_data(data, titles_from_data=True)
    categories = Reference(ws, min_col=1, min_row=min_row + 1, max_row=max_row)
    chart.set_categories(categories)
    chart.height = 8
    chart.width = 18
    ws.add_chart(chart, anchor)


def _add_bar_chart(ws, title, min_row, max_row, data_col, anchor):
    if max_row <= min_row:
        return
    chart = BarChart()
    chart.title = title
    chart.y_axis.title = "Count"
    data = Reference(ws, min_col=data_col, min_row=min_row, max_row=max_row)
    cats = Reference(ws, min_col=1, min_row=min_row + 1, max_row=max_row)
    chart.add_data(data, titles_from_data=True)
    chart.set_categories(cats)
    chart.height = 8
    chart.width = 14
    ws.add_chart(chart, anchor)


def _add_pie_chart(ws, title, min_row, max_row, data_col, anchor):
    if max_row <= min_row:
        return
    chart = PieChart()
    chart.title = title
    data = Reference(ws, min_col=data_col, min_row=min_row, max_row=max_row)
    labels = Reference(ws, min_col=1, min_row=min_row + 1, max_row=max_row)
    chart.add_data(data, titles_from_data=True)
    chart.set_categories(labels)
    chart.height = 8
    chart.width = 12
    ws.add_chart(chart, anchor)


def build_monthly_analytics_excel(office=None):
    office_keys = _normalize_office_selection(office)
    payload = build_monthly_analytics_payload(office_keys)
    cleanup_summary = _latest_duplicate_cleanup_summary()
    selected_labels = [payload[key]["label"] for key in office_keys]
    summary_subject = " and ".join(selected_labels)
    wb = Workbook()
    wb.properties.title = "NDOH Workforce Monthly Analytics Report"
    wb.properties.subject = f"{summary_subject} analytics"
    wb.properties.creator = "NDOH Workforce Online Registration System"
    guide = wb.active
    guide.title = "Read Me First"
    _style_sheet(guide, GOV_GOLD)
    _add_workbook_banner(guide, f"How To Read The Monthly Analytics Report for {summary_subject}")
    _append_table(
        guide,
        "Workbook Navigation",
        ["Step", "Sheet", "What It Contains", "How To Use It"],
        _build_workbook_navigation_rows(),
        freeze=False,
        auto_filter=False,
    )
    _append_table(
        guide,
        "How To Read The Monthly Analytics Workbook",
        ["Section", "What It Means", "How To Explain It"],
        _build_guide_rows(),
        freeze=False,
        auto_filter=True,
    )
    guide.freeze_panes = "A2"

    summary = wb.create_sheet("Executive Summary")
    summary.sheet_properties.tabColor = GOV_NAVY
    _add_workbook_banner(summary, f"Monthly Analytics Report for {summary_subject}")
    summary_rows = [
        ["Report Owner", "National Department of Health"],
        ["System", "Workforce Online Registration System"],
        ["Generated On", payload["combined"]["generated_on"]],
    ]
    for office_key in office_keys:
        summary_rows.append([f"{payload[office_key]['label']} Current Registry People", payload[office_key]["current_registry_total"]])
        summary_rows.append([f"{payload[office_key]['label']} Imported Record Rows", payload[office_key]["total_records"]])
        if payload[office_key]["latest_batch_row"]:
            summary_rows.append([f"{payload[office_key]['label']} Latest Source File", payload[office_key]["latest_batch_row"]["source_file_name"]])
            summary_rows.append([f"{payload[office_key]['label']} Latest Completed Import", payload[office_key]["latest_batch_row"]["completed_at"]])
            summary_rows.append([f"{payload[office_key]['label']} Recent Source Sheets", _source_sheet_summary(payload[office_key]["source_sheet_rows"])])
    if len(office_keys) > 1:
        summary_rows.extend([
            ["Duplicate Reviews", payload["combined"]["duplicate_reviews"]],
            ["Open Missing Data Reviews", payload["combined"]["open_missing_reviews"]],
        ])
    _append_table(summary, "NDOH Workforce Monthly Analytics Report", ["Metric", "Value"], summary_rows)

    for office_key in office_keys:
        data = payload[office_key]
        overview_ws = wb.create_sheet(f"{data['label']} Overview"[:31])
        _style_sheet(overview_ws, GOV_NAVY if office_key == "medical" else GOV_GREEN)
        _append_table(
            overview_ws,
            f"{data['label']} Current Registry Totals",
            ["Category", "Current Registry Count", "Active Count", "Latest Record Update", "Source Table"],
            [
                [row["label"], row["current_count"], row["active_count"], row["latest_updated"], row["source_model"]]
                for row in data["current_registry_rows"]
            ],
            freeze=False,
            auto_filter=False,
        )
        _append_table(
            overview_ws,
            f"{data['label']} Comparison Between Live Registry and Imported Rows",
            ["Measure", "Value", "Meaning"],
            _build_comparison_rows(data),
            freeze=False,
            auto_filter=False,
        )
        _append_table(
            overview_ws,
            f"{data['label']} How Current Totals Were Calculated",
            [
                "Category",
                "Current Count",
                "Active Count",
                "Source Table",
                "Calculation Used",
                "Latest Live Update",
                "Latest Source Import",
                "Source File",
                "Recent Source Sheets",
            ],
            _build_registry_explanation_rows(data),
            freeze=False,
            auto_filter=False,
        )
        if data["latest_batch_row"]:
            _append_table(
                overview_ws,
                f"{data['label']} Latest Data Source",
                ["Metric", "Value"],
                [
                    ["Latest Batch ID", data["latest_batch_row"]["id"]],
                    ["Source File", data["latest_batch_row"]["source_file_name"]],
                    ["Source Kind", data["latest_batch_row"]["source_kind"]],
                    ["Batch Status", data["latest_batch_row"]["status"]],
                    ["Processed Rows", data["latest_batch_row"]["processed_rows"]],
                    ["Total Rows", data["latest_batch_row"]["total_rows"]],
                    ["Completed At", data["latest_batch_row"]["completed_at"]],
                    ["Started At", data["latest_batch_row"]["started_at"]],
                ],
                freeze=False,
                auto_filter=False,
            )
            _append_table(
                overview_ws,
                f"{data['label']} Latest Source Sheets",
                ["Sheet Name", "Sheet Type", "Status", "Imported Rows", "Skipped Rows"],
                [
                    [
                        row["sheet_name"],
                        row["sheet_type"],
                        row["status"],
                        row["imported_rows"],
                        row["skipped_rows"],
                    ]
                    for row in data["source_sheet_rows"]
                ],
                freeze=False,
                auto_filter=False,
            )
        _append_table(
            overview_ws,
            f"{data['label']} Management Questions and Answers",
            ["Question", "Answer", "Evidence / Where To Check"],
            _build_management_answer_rows(data),
            freeze=False,
            auto_filter=False,
        )
        if office_key == "nursing":
            _append_table(
                overview_ws,
                "Nursing Council Data Cleansing and Corrections",
                ["Measure", "Current Value", "Explanation"],
                _build_cleansing_rows(data, cleanup_summary),
                freeze=False,
                auto_filter=False,
            )
            _append_table(
                overview_ws,
                "Nursing Council Corrected Current Statistics",
                ["Measure", "Current Total", "How It Was Calculated"],
                _build_nursing_correction_rows(data),
                freeze=False,
                auto_filter=False,
            )
        overview_ws.freeze_panes = "A2"
        overview_ws.auto_filter.ref = None
        overview_ws.sheet_view.topLeftCell = "A1"
        overview_ws.column_dimensions["A"].width = 34
        overview_ws.column_dimensions["B"].width = 22
        overview_ws.column_dimensions["C"].width = 42
        overview_ws.column_dimensions["D"].width = 22
        overview_ws.column_dimensions["E"].width = 24
        overview_ws.column_dimensions["F"].width = 20
        overview_ws.column_dimensions["G"].width = 20
        overview_ws.column_dimensions["H"].width = 28
        overview_ws.column_dimensions["I"].width = 34

        ws = wb.create_sheet(f"{data['label']} Monthly"[:31])
        _style_sheet(ws, GOV_GREEN if office_key == "nursing" else GOV_NAVY)
        monthly_rows = [
            [
                row["month"],
                row["full"],
                row["practicing_license"],
                row["provisional"],
                row["temporary"],
                row["workforce_listing"],
                row["payment"],
                row["unique_people"],
                row["total_records"],
            ]
            for row in data["monthly_rows"]
        ]
        start, end = _append_table(
            ws,
            f"{data['label']} Monthly Analytics",
            ["Month", "Full Registration", "Practising Licence", "Provisional", "Temporary", "Workforce Listing", "Payment", "Unique People", "Total Records"],
            monthly_rows,
        )
        _add_line_chart(ws, f"{data['label']} Monthly Trend", start, end, [2, 3, 8], "K3")

        yearly_ws = wb.create_sheet(f"{data['label']} Yearly"[:31])
        _style_sheet(yearly_ws, GOV_GOLD if office_key == "nursing" else GOV_NAVY)
        yearly_rows = [
            [row["year"], row["full"], row["practicing_license"], row["provisional"], row["temporary"], row["workforce_listing"], row["payment"], row["total_records"]]
            for row in data["yearly_rows"]
        ]
        y_start, y_end = _append_table(
            yearly_ws,
            f"{data['label']} Yearly Analytics",
            ["Year", "Full Registration", "Practising Licence", "Provisional", "Temporary", "Workforce Listing", "Payment", "Total Records"],
            yearly_rows,
        )
        _add_bar_chart(yearly_ws, f"{data['label']} Total Records by Year", y_start, y_end, 8, "J3")

        mix_ws = wb.create_sheet(f"{data['label']} Mix"[:31])
        _style_sheet(mix_ws, GOV_GREEN if office_key == "nursing" else GOV_GOLD)
        r_start, r_end = _append_table(
            mix_ws,
            f"{data['label']} Record Type Mix",
            ["Record Type", "Count"],
            [[row["record_type"], row["count"]] for row in data["record_type_rows"]],
        )
        _add_pie_chart(mix_ws, f"{data['label']} Record Types", r_start, r_end, 2, "D3")
        c_start, c_end = _append_table(
            mix_ws,
            f"{data['label']} Top Categories",
            ["Category", "Count"],
            data["category_rows"],
        )
        _add_bar_chart(mix_ws, f"{data['label']} Categories", c_start, c_end, 2, "D20")

        app_ws = wb.create_sheet(f"{data['label']} Admin"[:31])
        _style_sheet(app_ws, GOV_NAVY)
        _append_table(
            app_ws,
            f"{data['label']} Applications by Status",
            ["Status", "Count"],
            [[row["status"], row["count"]] for row in data["application_rows"]],
        )
        _append_table(
            app_ws,
            f"{data['label']} Recent Import Batches",
            ["Batch ID", "File", "Kind", "Status", "Processed Rows", "Total Rows", "Completed At"],
            [
                [
                    row["id"],
                    row["source_file_name"],
                    row["source_kind"],
                    row["status"],
                    row["processed_rows"],
                    row["total_rows"],
                    timezone.localtime(row["completed_at"]).strftime("%Y-%m-%d %H:%M") if row["completed_at"] else "",
                ]
                for row in data["batch_rows"]
            ],
        )

    for worksheet in wb.worksheets:
        _style_sheet(worksheet, worksheet.sheet_properties.tabColor.rgb if worksheet.sheet_properties.tabColor and worksheet.sheet_properties.tabColor.type == "rgb" else GOV_GREEN)
    wb.active = 0

    output = BytesIO()
    wb.save(output)
    return output.getvalue()


def _pdf_table(headers, rows, available_width):
    styles = getSampleStyleSheet()
    header_style = ParagraphStyle("Header", parent=styles["Normal"], fontSize=7, leading=8, fontName="Helvetica-Bold", textColor=colors.white, alignment=1)
    cell_style = ParagraphStyle("Cell", parent=styles["Normal"], fontSize=7, leading=8, splitLongWords=True)
    data = [[Paragraph(str(header), header_style) for header in headers]]
    for row in rows:
        data.append([Paragraph(str(value or ""), cell_style) for value in row])
    widths = [available_width / len(headers)] * len(headers)
    table = Table(data, repeatRows=1, colWidths=widths)
    table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor(f"#{GOV_GREEN}")),
        ("GRID", (0, 0), (-1, -1), 0.25, colors.HexColor(f"#{GOV_BORDER}")),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor(f"#{GOV_LIGHT}")]),
        ("LEFTPADDING", (0, 0), (-1, -1), 3),
        ("RIGHTPADDING", (0, 0), (-1, -1), 3),
    ]))
    return table


def _pdf_key_value_table(rows, available_width):
    return _pdf_table(["Metric", "Value"], rows, available_width)


def _draw_government_pdf_frame(canvas, doc):
    canvas.saveState()
    width, height = landscape(letter)
    canvas.setFillColor(colors.HexColor(f"#{GOV_NAVY}"))
    canvas.rect(0, height - 34, width, 34, stroke=0, fill=1)
    canvas.setFillColor(colors.HexColor(f"#{GOV_GOLD}"))
    canvas.rect(0, height - 38, width, 4, stroke=0, fill=1)
    canvas.setFillColor(colors.white)
    canvas.setFont("Helvetica-Bold", 9)
    logo_path = _ndoh_logo_path()
    text_x = 24
    if logo_path:
        try:
            canvas.drawImage(str(logo_path), 24, height - 31, width=24, height=24, preserveAspectRatio=True, mask="auto")
            text_x = 56
        except Exception:
            text_x = 24
    canvas.drawString(text_x, height - 21, "National Department of Health | Workforce Online Registration System")
    canvas.setFillColor(colors.HexColor(f"#{GOV_NAVY}"))
    canvas.setFont("Helvetica", 7)
    canvas.drawString(24, 14, "Confidential registry analytics - generated for official NDOH reporting use")
    canvas.drawRightString(width - 24, 14, f"Page {doc.page}")
    canvas.restoreState()


def build_monthly_analytics_pdf(office=None):
    office_keys = _normalize_office_selection(office)
    payload = build_monthly_analytics_payload(office_keys)
    cleanup_summary = _latest_duplicate_cleanup_summary()
    selected_labels = [payload[key]["label"] for key in office_keys]
    summary_subject = " and ".join(selected_labels)
    output = BytesIO()
    doc = SimpleDocTemplate(output, pagesize=landscape(letter), leftMargin=24, rightMargin=24, topMargin=52, bottomMargin=28)
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle("ReportTitle", parent=styles["Heading1"], fontSize=16, leading=19, textColor=colors.HexColor(f"#{GOV_NAVY}"), alignment=1, spaceAfter=4)
    subtitle_style = ParagraphStyle("ReportSubtitle", parent=styles["Normal"], fontSize=9, leading=11, textColor=colors.HexColor(f"#{GOV_GOLD}"), alignment=1, spaceAfter=8)
    section_style = ParagraphStyle("SectionTitle", parent=styles["Heading2"], fontSize=11, leading=13, textColor=colors.HexColor(f"#{GOV_GREEN}"), spaceBefore=8)
    normal = ParagraphStyle("ReportNormal", parent=styles["Normal"], fontSize=8, leading=10)
    story = [
        Paragraph("National Department of Health", title_style),
        Paragraph(f"{summary_subject} Workforce Monthly Analytics Report", subtitle_style),
        Paragraph(f"Generated on {payload['combined']['generated_on']}. This report covers live registry totals, imported records, applications, and import activity for official registry monitoring.", normal),
        Spacer(1, 8),
    ]
    summary_rows = [[f"{payload[key]['label']} Imported Records", payload[key]["total_records"]] for key in office_keys]
    summary_rows = []
    for key in office_keys:
        summary_rows.extend([
            [f"{payload[key]['label']} Current Registry People", payload[key]["current_registry_total"]],
            [f"{payload[key]['label']} Imported Record Rows", payload[key]["total_records"]],
        ])
        if payload[key]["latest_batch_row"]:
            summary_rows.append([f"{payload[key]['label']} Latest Source File", payload[key]["latest_batch_row"]["source_file_name"]])
            summary_rows.append([f"{payload[key]['label']} Latest Completed Import", payload[key]["latest_batch_row"]["completed_at"]])
            summary_rows.append([f"{payload[key]['label']} Recent Source Sheets", _source_sheet_summary(payload[key]["source_sheet_rows"])])
    if len(office_keys) > 1:
        summary_rows.extend([
            ["Duplicate Reviews", payload["combined"]["duplicate_reviews"]],
            ["Open Missing Data Reviews", payload["combined"]["open_missing_reviews"]],
        ])
    story.append(_pdf_table(["Metric", "Value"], summary_rows, doc.width))
    story.extend([
        Spacer(1, 8),
        Paragraph("How To Understand These Figures", section_style),
        _pdf_table(["Section", "Meaning", "How To Explain It"], _build_guide_rows(), doc.width),
    ])
    for office_key in office_keys:
        data = payload[office_key]
        story.extend([
            Spacer(1, 10),
            Paragraph(data["label"], section_style),
            _pdf_table(
                ["Category", "Current Registry Count", "Active Count", "Latest Record Update", "Source Table"],
                [
                    [row["label"], row["current_count"], row["active_count"], row["latest_updated"], row["source_model"]]
                    for row in data["current_registry_rows"]
                ],
                doc.width,
            ),
            Spacer(1, 6),
            _pdf_table(
                ["Measure", "Value", "Meaning"],
                _build_comparison_rows(data),
                doc.width,
            ),
            Spacer(1, 6),
            _pdf_table(
                [
                    "Category",
                    "Current Count",
                    "Active Count",
                    "Source Table",
                    "Calculation Used",
                    "Latest Live Update",
                    "Latest Source Import",
                    "Source File",
                    "Recent Source Sheets",
                ],
                _build_registry_explanation_rows(data),
                doc.width,
            ),
            Spacer(1, 6),
        ])
        if office_key == "nursing":
            story.extend([
                Paragraph("Nursing Council Corrections and Data Cleansing Notes", section_style),
                _pdf_table(
                    ["Measure", "Current Value", "Explanation"],
                    _build_cleansing_rows(data, cleanup_summary),
                    doc.width,
                ),
                Spacer(1, 6),
                _pdf_table(
                    ["Measure", "Current Total", "How It Was Calculated"],
                    _build_nursing_correction_rows(data),
                    doc.width,
                ),
                Spacer(1, 6),
            ])
        if data["latest_batch_row"]:
            story.extend([
                _pdf_key_value_table(
                    [
                        ["Latest Batch ID", data["latest_batch_row"]["id"]],
                        ["Source File", data["latest_batch_row"]["source_file_name"]],
                        ["Source Kind", data["latest_batch_row"]["source_kind"]],
                        ["Completed At", data["latest_batch_row"]["completed_at"]],
                        ["Processed Rows", data["latest_batch_row"]["processed_rows"]],
                        ["Total Rows", data["latest_batch_row"]["total_rows"]],
                    ],
                    doc.width * 0.58,
                ),
                Spacer(1, 6),
            ])
            if data["source_sheet_rows"]:
                story.extend([
                    _pdf_table(
                        ["Sheet Name", "Sheet Type", "Status", "Imported Rows", "Skipped Rows"],
                        [
                            [
                                row["sheet_name"],
                                row["sheet_type"],
                                row["status"],
                                row["imported_rows"],
                                row["skipped_rows"],
                            ]
                            for row in data["source_sheet_rows"][:12]
                        ],
                        doc.width,
                    ),
                    Spacer(1, 6),
                ])
        story.extend([
            _pdf_table(
                ["Question", "Answer", "Evidence / Where To Check"],
                _build_management_answer_rows(data),
                doc.width,
            ),
            Spacer(1, 6),
            _pdf_table(
                ["Month", "Full", "Practising", "Provisional", "Temporary", "Workforce", "Payments", "Unique", "Total"],
                [
                    [
                        row["month"],
                        row["full"],
                        row["practicing_license"],
                        row["provisional"],
                        row["temporary"],
                        row["workforce_listing"],
                        row["payment"],
                        row["unique_people"],
                        row["total_records"],
                    ]
                    for row in data["monthly_rows"][-12:]
                ],
                doc.width,
            ),
            Spacer(1, 6),
            _pdf_table(
                ["Record Type", "Count"],
                [[row["record_type"], row["count"]] for row in data["record_type_rows"]],
                doc.width * 0.45,
            ),
        ])
    doc.build(story, onFirstPage=_draw_government_pdf_frame, onLaterPages=_draw_government_pdf_frame)
    return output.getvalue()


def build_yearly_analytics_excel(office=None):
    office_keys = _normalize_office_selection(office)
    payload = build_monthly_analytics_payload(office_keys)
    selected_labels = [payload[key]["label"] for key in office_keys]
    summary_subject = " and ".join(selected_labels)
    wb = Workbook()
    wb.properties.title = "NDOH Workforce Yearly Analytics Report"
    wb.properties.subject = f"{summary_subject} yearly analytics"
    wb.properties.creator = "NDOH Workforce Online Registration System"

    summary = wb.active
    summary.title = "Executive Summary"
    summary.sheet_properties.tabColor = GOV_NAVY
    _add_workbook_banner(summary, f"Yearly Analytics Report for {summary_subject}")
    summary_rows = [
        ["Report Owner", "National Department of Health"],
        ["System", "Workforce Online Registration System"],
        ["Generated On", payload["combined"]["generated_on"]],
    ]
    for office_key in office_keys:
        office_data = payload[office_key]
        summary_rows.extend([
            [f"{office_data['label']} Current Registry People", office_data["current_registry_total"]],
            [f"{office_data['label']} Imported Record Rows", office_data["total_records"]],
        ])
        if office_data["latest_batch_row"]:
            summary_rows.append([f"{office_data['label']} Latest Source File", office_data["latest_batch_row"]["source_file_name"]])
            summary_rows.append([f"{office_data['label']} Latest Completed Import", office_data["latest_batch_row"]["completed_at"]])
    _append_table(summary, "NDOH Workforce Yearly Analytics Report", ["Metric", "Value"], summary_rows)

    guide = wb.create_sheet("How To Use")
    _style_sheet(guide, GOV_GOLD)
    _append_table(
        guide,
        "How To Read The Yearly Analytics Workbook",
        ["Section", "What It Means", "How To Use It"],
        [
            ["Current Registry Totals", "Live people counts held in the main registry tables.", "Use this when management asks how many practitioners are currently in the system."],
            ["Yearly Analytics", "Historical record counts grouped by record year and record type.", "Use this to explain activity patterns across years rather than by month."],
            ["Application Status", "Current application totals by status for the selected office scope.", "Use this to explain the live review workload."],
            ["Source Batches", "Latest and recent import batches used to refresh registry statistics.", "Use this when leadership asks where the latest figures came from."],
        ],
    )

    for office_key in office_keys:
        data = payload[office_key]
        overview_ws = wb.create_sheet(f"{data['label']} Overview"[:31])
        _style_sheet(overview_ws, GOV_GREEN if office_key == "nursing" else GOV_NAVY)
        _append_table(
            overview_ws,
            f"{data['label']} Current Registry Totals",
            ["Category", "Current Registry Count", "Active Count", "Latest Record Update", "Source Table"],
            [
                [row["label"], row["current_count"], row["active_count"], row["latest_updated"], row["source_model"]]
                for row in data["current_registry_rows"]
            ],
        )
        if data["latest_batch_row"]:
            _append_table(
                overview_ws,
                f"{data['label']} Latest Data Source",
                ["Metric", "Value"],
                [
                    ["Latest Batch ID", data["latest_batch_row"]["id"]],
                    ["Source File", data["latest_batch_row"]["source_file_name"]],
                    ["Source Kind", data["latest_batch_row"]["source_kind"]],
                    ["Batch Status", data["latest_batch_row"]["status"]],
                    ["Processed Rows", data["latest_batch_row"]["processed_rows"]],
                    ["Total Rows", data["latest_batch_row"]["total_rows"]],
                    ["Completed At", data["latest_batch_row"]["completed_at"]],
                ],
            )

        yearly_ws = wb.create_sheet(f"{data['label']} Yearly"[:31])
        _style_sheet(yearly_ws, GOV_GOLD if office_key == "nursing" else GOV_NAVY)
        yearly_rows = [
            [row["year"], row["full"], row["practicing_license"], row["provisional"], row["temporary"], row["workforce_listing"], row["payment"], row["total_records"]]
            for row in data["yearly_rows"]
        ]
        start, end = _append_table(
            yearly_ws,
            f"{data['label']} Yearly Analytics",
            ["Year", "Full Registration", "Practising Licence", "Provisional", "Temporary", "Workforce Listing", "Payment", "Total Records"],
            yearly_rows,
        )
        _add_bar_chart(yearly_ws, f"{data['label']} Total Records by Year", start, end, 8, "J3")

        admin_ws = wb.create_sheet(f"{data['label']} Admin"[:31])
        _style_sheet(admin_ws, GOV_NAVY)
        _append_table(
            admin_ws,
            f"{data['label']} Applications by Status",
            ["Status", "Count"],
            [[row["status"], row["count"]] for row in data["application_rows"]],
        )
        _append_table(
            admin_ws,
            f"{data['label']} Record Type Mix",
            ["Record Type", "Count"],
            [[row["record_type"], row["count"]] for row in data["record_type_rows"]],
        )
        _append_table(
            admin_ws,
            f"{data['label']} Recent Import Batches",
            ["Batch ID", "File", "Kind", "Status", "Processed Rows", "Total Rows", "Completed At"],
            [
                [
                    row["id"],
                    row["source_file_name"],
                    row["source_kind"],
                    row["status"],
                    row["processed_rows"],
                    row["total_rows"],
                    timezone.localtime(row["completed_at"]).strftime("%Y-%m-%d %H:%M") if row["completed_at"] else "",
                ]
                for row in data["batch_rows"]
            ],
        )

    output = BytesIO()
    wb.save(output)
    return output.getvalue()


def build_yearly_analytics_pdf(office=None):
    office_keys = _normalize_office_selection(office)
    payload = build_monthly_analytics_payload(office_keys)
    selected_labels = [payload[key]["label"] for key in office_keys]
    summary_subject = " and ".join(selected_labels)
    output = BytesIO()
    doc = SimpleDocTemplate(output, pagesize=landscape(letter), leftMargin=24, rightMargin=24, topMargin=52, bottomMargin=28)
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle("ReportTitle", parent=styles["Heading1"], fontSize=16, leading=19, textColor=colors.HexColor(f"#{GOV_NAVY}"), alignment=1, spaceAfter=4)
    subtitle_style = ParagraphStyle("ReportSubtitle", parent=styles["Normal"], fontSize=9, leading=11, textColor=colors.HexColor(f"#{GOV_GOLD}"), alignment=1, spaceAfter=8)
    section_style = ParagraphStyle("SectionTitle", parent=styles["Heading2"], fontSize=11, leading=13, textColor=colors.HexColor(f"#{GOV_GREEN}"), spaceBefore=8)
    normal = ParagraphStyle("ReportNormal", parent=styles["Normal"], fontSize=8, leading=10)

    story = [
        Paragraph("National Department of Health", title_style),
        Paragraph(f"{summary_subject} Workforce Yearly Analytics Report", subtitle_style),
        Paragraph(f"Generated on {payload['combined']['generated_on']}. This report focuses on live registry totals, yearly historical activity, applications, and source import status.", normal),
        Spacer(1, 8),
    ]
    summary_rows = []
    for key in office_keys:
        summary_rows.extend([
            [f"{payload[key]['label']} Current Registry People", payload[key]["current_registry_total"]],
            [f"{payload[key]['label']} Imported Record Rows", payload[key]["total_records"]],
        ])
        if payload[key]["latest_batch_row"]:
            summary_rows.append([f"{payload[key]['label']} Latest Source File", payload[key]["latest_batch_row"]["source_file_name"]])
            summary_rows.append([f"{payload[key]['label']} Latest Completed Import", payload[key]["latest_batch_row"]["completed_at"]])
    story.append(_pdf_table(["Metric", "Value"], summary_rows, doc.width))

    for office_key in office_keys:
        data = payload[office_key]
        story.extend([
            Spacer(1, 10),
            Paragraph(data["label"], section_style),
            _pdf_table(
                ["Category", "Current Registry Count", "Active Count", "Latest Record Update", "Source Table"],
                [
                    [row["label"], row["current_count"], row["active_count"], row["latest_updated"], row["source_model"]]
                    for row in data["current_registry_rows"]
                ],
                doc.width,
            ),
            Spacer(1, 6),
        ])
        if data["latest_batch_row"]:
            story.extend([
                _pdf_key_value_table(
                    [
                        ["Latest Batch ID", data["latest_batch_row"]["id"]],
                        ["Source File", data["latest_batch_row"]["source_file_name"]],
                        ["Source Kind", data["latest_batch_row"]["source_kind"]],
                        ["Completed At", data["latest_batch_row"]["completed_at"]],
                        ["Processed Rows", data["latest_batch_row"]["processed_rows"]],
                        ["Total Rows", data["latest_batch_row"]["total_rows"]],
                    ],
                    doc.width * 0.58,
                ),
                Spacer(1, 6),
            ])

        story.extend([
            _pdf_table(
                ["Year", "Full", "Practising", "Provisional", "Temporary", "Workforce", "Payments", "Total"],
                [
                    [
                        row["year"],
                        row["full"],
                        row["practicing_license"],
                        row["provisional"],
                        row["temporary"],
                        row["workforce_listing"],
                        row["payment"],
                        row["total_records"],
                    ]
                    for row in data["yearly_rows"]
                ],
                doc.width,
            ),
            Spacer(1, 6),
            _pdf_table(
                ["Application Status", "Count"],
                [[row["status"], row["count"]] for row in data["application_rows"]],
                doc.width * 0.45,
            ),
        ])

    doc.build(story, onFirstPage=_draw_government_pdf_frame, onLaterPages=_draw_government_pdf_frame)
    return output.getvalue()


FINANCIAL_CATEGORY_LABELS = {
    "renewal": "Renewals / Practising Licences",
    "full_registration": "Full Registration",
    "new_application": "New Applications / Provisional",
    "temporary": "Temporary / Special Cases",
    "other": "Other Receipts",
}
FINANCIAL_CATEGORY_ORDER = [
    "renewal",
    "full_registration",
    "new_application",
    "temporary",
    "other",
]
FINANCIAL_FORM_CATEGORY_MAP = {
    "nursing": {
        "renewal": {"NC3"},
        "full_registration": {"NC2", "NC5", "NC6", "NC7", "NC10", "NC11"},
        "new_application": {"G1", "G2", "G3", "G4", "G5", "G6", "G7", "NC1", "NC4"},
        "temporary": {"NC8", "NC9"},
    },
    "medical": {
        "renewal": {"MD2", "MBRN"},
        "full_registration": {"MD1", "CHW1"},
        "new_application": {"MBSP", "MBAC", "MBPF", "MBTC"},
        "temporary": set(),
    },
}


def _financial_office_scope(office):
    config = _office_config(office)
    return {
        "label": config["label"],
        "batch_kinds": config["batch_kinds"],
        "forms": config["forms"],
    }


def _financial_decimal(value):
    return Decimal(str(value or 0))


def _financial_month_key(dt_value):
    return dt_value.strftime("%Y-%m")


def _financial_year_key(dt_value):
    return dt_value.year


def _coerce_financial_date(value):
    if not value:
        return None
    if isinstance(value, datetime):
        if timezone.is_aware(value):
            value = timezone.localtime(value)
        return value.date()
    if isinstance(value, date):
        return value
    return None


def _financial_date_reason(value, today):
    if not value:
        return ""
    if value > today:
        return f"future date after current system date {today.isoformat()}"
    if value < FINANCIAL_MIN_VALID_DATE:
        return f"date before valid registry reporting range {FINANCIAL_MIN_VALID_DATE.isoformat()}"
    return ""


def _resolve_financial_reporting_date(source_date, fallback_date, today, issues, issue_context):
    raw_date = _coerce_financial_date(source_date)
    fallback = _coerce_financial_date(fallback_date) or today
    if fallback > today or fallback < FINANCIAL_MIN_VALID_DATE:
        fallback = today

    reason = _financial_date_reason(raw_date, today)
    if raw_date and not reason:
        return {
            "period_date": raw_date,
            "display_date": raw_date.isoformat(),
            "date_note": "",
            "date_flagged": False,
            "source_date": raw_date,
        }

    if raw_date and reason:
        issue = {
            "source": issue_context.get("source", ""),
            "source_date": raw_date.isoformat(),
            "reporting_date": fallback.isoformat(),
            "reason": reason,
            "reference": issue_context.get("reference") or "-",
            "amount": issue_context.get("amount") or Decimal("0"),
            "detail": issue_context.get("detail") or "",
            "record_id": issue_context.get("record_id") or "",
        }
        issues.append(issue)
        return {
            "period_date": fallback,
            "display_date": fallback.isoformat(),
            "date_note": f"Flagged source date {raw_date.isoformat()}: {reason}. Grouped under {fallback.isoformat()}.",
            "date_flagged": True,
            "source_date": raw_date,
        }

    return {
        "period_date": fallback,
        "display_date": fallback.isoformat(),
        "date_note": "No receipt/payment date supplied. Grouped under transaction/import date.",
        "date_flagged": False,
        "source_date": None,
    }


def _manual_receipt_office(receipt):
    if receipt.application_id and receipt.application:
        return "medical" if receipt.application.form_code in MEDICAL_FORMS else "nursing"
    if receipt.user_id and receipt.user:
        if receipt.user.role in {"doctor", "chw"}:
            return "medical"
        if receipt.user.role in {"nurse", "nurse_aide", "graduand"}:
            return "nursing"
        if receipt.user.role == "registrar":
            department = (receipt.user.department or "").lower()
            return "medical" if "medical" in department else "nursing"
    return None


def _is_imported_receipt(receipt):
    payment_method = (receipt.payment_method or "").lower()
    description = (receipt.description or "").lower()
    return "imported" in payment_method or description.startswith("imported payment row")


def _manual_receipt_category(receipt, office):
    form_code = ""
    if receipt.application_id and receipt.application:
        form_code = (receipt.application.form_code or "").upper()
    if form_code:
        for category, form_codes in FINANCIAL_FORM_CATEGORY_MAP[office].items():
            if form_code in form_codes:
                return category
    description = (receipt.description or "").lower()
    if "renew" in description or "practising" in description or "practicing" in description:
        return "renewal"
    if "full" in description:
        return "full_registration"
    if "provisional" in description or "application" in description:
        return "new_application"
    if "temporary" in description:
        return "temporary"
    return "other"


def _imported_receipt_category(record):
    if record.record_type in {"payment", "practicing_license"}:
        return "renewal"
    if record.record_type == "full":
        return "full_registration"
    if record.record_type == "provisional":
        return "new_application"
    if record.record_type == "temporary":
        return "temporary"
    return "other"


def _financial_init_period_row(period_label):
    row = {
        "period": period_label,
        "manual_amount": Decimal("0"),
        "imported_amount": Decimal("0"),
        "total_amount": Decimal("0"),
        "manual_count": 0,
        "imported_count": 0,
    }
    for key in FINANCIAL_CATEGORY_ORDER:
        row[f"{key}_amount"] = Decimal("0")
    return row


def _financial_add_period_entry(bucket, amount, category, source):
    amount = _financial_decimal(amount)
    bucket[f"{source}_amount"] += amount
    bucket["total_amount"] += amount
    bucket[f"{source}_count"] += 1
    bucket[f"{category}_amount"] += amount


def _financial_forecast_rows(monthly_rows, periods=6):
    actual_rows = [row for row in monthly_rows if row["total_amount"] > 0]
    recent_rows = actual_rows[-6:] if len(actual_rows) >= 6 else actual_rows
    if recent_rows:
        average = sum((row["total_amount"] for row in recent_rows), Decimal("0")) / Decimal(len(recent_rows))
    else:
        average = Decimal("0")

    forecast_rows = []
    if actual_rows:
        last_period = actual_rows[-1]["period"] + "-01"
        current = date.fromisoformat(last_period)
    else:
        today = date.today()
        current = date(today.year, today.month, 1)

    for _ in range(periods):
        next_month = current.month + 1
        next_year = current.year
        if next_month > 12:
            next_month = 1
            next_year += 1
        current = date(next_year, next_month, 1)
        forecast_rows.append({
            "period": current.strftime("%Y-%m"),
            "forecast_amount": average.quantize(Decimal("0.01")),
            "low_amount": (average * Decimal("0.85")).quantize(Decimal("0.01")),
            "high_amount": (average * Decimal("1.15")).quantize(Decimal("0.01")),
        })
    return forecast_rows


def _financial_summary_rows(office_data):
    return [
        ["Portal / Manual Receipts Amount", office_data["manual_completed_total"]],
        ["Imported Spreadsheet Receipt Amount", office_data["imported_total"]],
        ["Combined Financial Total", office_data["combined_total"]],
        ["Current Month Completed Manual Receipts", office_data["manual_current_month_total"]],
        ["Current Year Completed Manual Receipts", office_data["manual_current_year_total"]],
        ["Current Month Combined Receipts", office_data["combined_current_month_total"]],
        ["Current Year Combined Receipts", office_data["combined_current_year_total"]],
        ["Completed Manual Receipts Count", office_data["manual_completed_count"]],
        ["Imported Spreadsheet Rows Processed", office_data["imported_count"]],
        ["Pending Manual Receipts", office_data["manual_pending_count"]],
        ["Flagged Date Issues", office_data["date_quality_issue_count"]],
    ]


def build_financial_forecast_payload(office=None, generated_by=None):
    office_keys = _normalize_office_selection(office)
    payload = {
        "generated_on": timezone.localtime().strftime("%d %b %Y %H:%M"),
        "generated_by": generated_by or "System generated",
        "office_keys": office_keys,
        "offices": {},
    }

    all_receipts = list(Receipt.objects.select_related("user", "application").order_by("-receipt_date", "-transaction_date"))

    for office_key in office_keys:
        config = _financial_office_scope(office_key)
        monthly = defaultdict(lambda: _financial_init_period_row(""))
        yearly = defaultdict(lambda: _financial_init_period_row(""))
        category_totals = {key: Decimal("0") for key in FINANCIAL_CATEGORY_ORDER}
        recent_transactions = []

        manual_completed_total = Decimal("0")
        manual_pending_count = 0
        manual_completed_count = 0
        manual_current_month_total = Decimal("0")
        manual_current_year_total = Decimal("0")
        combined_current_month_total = Decimal("0")
        combined_current_year_total = Decimal("0")
        today = timezone.localdate()
        date_quality_issues = []

        scoped_receipts = []
        for receipt in all_receipts:
            if _manual_receipt_office(receipt) != office_key:
                continue
            if _is_imported_receipt(receipt):
                continue
            scoped_receipts.append(receipt)

        for receipt in scoped_receipts:
            if receipt.status != "completed":
                if receipt.status == "pending":
                    manual_pending_count += 1
                continue
            amount = _financial_decimal(receipt.amount)
            date_info = _resolve_financial_reporting_date(
                receipt.receipt_date,
                receipt.transaction_date,
                today,
                date_quality_issues,
                {
                    "source": "Portal Receipt",
                    "reference": receipt.official_receipt_no or receipt.receipt_number,
                    "amount": amount,
                    "detail": receipt.description or "Manual receipt",
                    "record_id": receipt.pk,
                },
            )
            period_date = date_info["period_date"]
            category = _manual_receipt_category(receipt, office_key)

            month_key = _financial_month_key(period_date)
            year_key = _financial_year_key(period_date)
            monthly_row = monthly[month_key]
            monthly_row["period"] = month_key
            yearly_row = yearly[year_key]
            yearly_row["period"] = year_key
            _financial_add_period_entry(monthly_row, amount, category, "manual")
            _financial_add_period_entry(yearly_row, amount, category, "manual")
            category_totals[category] += amount
            manual_completed_total += amount
            manual_completed_count += 1
            if period_date.year == today.year and period_date.month == today.month:
                manual_current_month_total += amount
                combined_current_month_total += amount
            if period_date.year == today.year:
                manual_current_year_total += amount
                combined_current_year_total += amount
            recent_transactions.append({
                "source": "Portal Receipt",
                "date": date_info["display_date"],
                "date_sort": period_date.isoformat(),
                "date_note": date_info["date_note"],
                "date_flagged": date_info["date_flagged"],
                "category": FINANCIAL_CATEGORY_LABELS[category],
                "reference": receipt.official_receipt_no or receipt.receipt_number,
                "amount": amount,
                "status": receipt.status.title(),
                "detail": receipt.description or (receipt.application.form_code if receipt.application_id and receipt.application else "Manual receipt"),
            })

        imported_records = list(
            PracticingLicenseRecord.objects.filter(
                batch__source_kind__in=config["batch_kinds"],
                amount__isnull=False,
            )
            .exclude(amount=0)
            .select_related("batch")
            .order_by("-payment_date", "-issued_date", "-created_at")
        )

        imported_total = Decimal("0")
        for record in imported_records:
            amount = _financial_decimal(record.amount)
            created = timezone.localtime(record.created_at) if timezone.is_aware(record.created_at) else record.created_at
            date_info = _resolve_financial_reporting_date(
                record.payment_date or record.issued_date,
                created,
                today,
                date_quality_issues,
                {
                    "source": "Spreadsheet Import",
                    "reference": record.reference_number or "-",
                    "amount": amount,
                    "detail": f"{record.source_sheet_name} row {record.source_row}",
                    "record_id": record.pk,
                },
            )
            period_date = date_info["period_date"]
            category = _imported_receipt_category(record)
            month_key = _financial_month_key(period_date)
            year_key = _financial_year_key(period_date)
            monthly_row = monthly[month_key]
            monthly_row["period"] = month_key
            yearly_row = yearly[year_key]
            yearly_row["period"] = year_key
            _financial_add_period_entry(monthly_row, amount, category, "imported")
            _financial_add_period_entry(yearly_row, amount, category, "imported")
            category_totals[category] += amount
            imported_total += amount
            if period_date.year == today.year and period_date.month == today.month:
                combined_current_month_total += amount
            if period_date.year == today.year:
                combined_current_year_total += amount
            recent_transactions.append({
                "source": "Spreadsheet Import",
                "date": date_info["display_date"],
                "date_sort": period_date.isoformat(),
                "date_note": date_info["date_note"],
                "date_flagged": date_info["date_flagged"],
                "category": FINANCIAL_CATEGORY_LABELS[category],
                "reference": record.reference_number or "-",
                "amount": amount,
                "status": record.get_record_type_display(),
                "detail": record.source_sheet_name,
            })

        monthly_rows = [dict(row) for _, row in sorted(monthly.items(), key=lambda item: item[0])]
        yearly_rows = [dict(row) for _, row in sorted(yearly.items(), key=lambda item: item[0])]
        forecast_rows = _financial_forecast_rows(monthly_rows)

        latest_completed_batches = list(
            DataImportBatch.objects.filter(source_kind__in=config["batch_kinds"], status="completed")
            .order_by("-completed_at", "-started_at")[:5]
        )

        office_data = {
            "label": config["label"],
            "manual_completed_total": manual_completed_total.quantize(Decimal("0.01")),
            "manual_completed_count": manual_completed_count,
            "manual_pending_count": manual_pending_count,
            "imported_total": imported_total.quantize(Decimal("0.01")),
            "imported_count": len(imported_records),
            "combined_total": (manual_completed_total + imported_total).quantize(Decimal("0.01")),
            "manual_current_month_total": manual_current_month_total.quantize(Decimal("0.01")),
            "manual_current_year_total": manual_current_year_total.quantize(Decimal("0.01")),
            "combined_current_month_total": combined_current_month_total.quantize(Decimal("0.01")),
            "combined_current_year_total": combined_current_year_total.quantize(Decimal("0.01")),
            "monthly_rows": monthly_rows,
            "yearly_rows": yearly_rows,
            "forecast_rows": forecast_rows,
            "date_quality_issues": sorted(date_quality_issues, key=lambda row: row["source_date"], reverse=True)[:50],
            "date_quality_issue_count": len(date_quality_issues),
            "category_rows": [
                {
                    "label": FINANCIAL_CATEGORY_LABELS[key],
                    "amount": category_totals[key].quantize(Decimal("0.01")),
                }
                for key in FINANCIAL_CATEGORY_ORDER
            ],
            "recent_transactions": sorted(recent_transactions, key=lambda row: row["date_sort"], reverse=True)[:20],
            "latest_batches": latest_completed_batches,
            "guide_rows": [
                "Portal / Manual Receipts come from the live Receipt table and represent recorded office or online receipts.",
                "Spreadsheet Receipt Amount is the PGK value imported from workbook payment rows. Spreadsheet Rows Processed is the count of imported receipt/payment rows.",
                "Combined totals join both streams for financial tracking, while keeping the two sources visible separately.",
                "Forecast values are based on the average of the latest six valid months with actual receipt amounts. Future source dates are excluded from driving the forecast and shown as date-quality issues.",
            ],
        }
        payload["offices"][office_key] = office_data

    if len(office_keys) > 1:
        payload["combined"] = {
            "manual_completed_total": sum((payload["offices"][key]["manual_completed_total"] for key in office_keys), Decimal("0.00")),
            "imported_total": sum((payload["offices"][key]["imported_total"] for key in office_keys), Decimal("0.00")),
            "combined_total": sum((payload["offices"][key]["combined_total"] for key in office_keys), Decimal("0.00")),
            "manual_pending_count": sum((payload["offices"][key]["manual_pending_count"] for key in office_keys), 0),
        }
    return payload


def build_financial_forecast_excel(office=None, generated_by=None):
    payload = build_financial_forecast_payload(office, generated_by=generated_by)
    wb = Workbook()
    wb.properties.title = "NDOH Financial Forecast Report"
    wb.properties.subject = "Regulatory bodies financial forecast"
    wb.properties.creator = "NDOH Workforce Online Registration System"

    summary = wb.active
    summary.title = "Executive Summary"
    summary.sheet_properties.tabColor = GOV_NAVY
    _add_workbook_banner(summary, "Financial Forecast and Receipt Tracking Report")
    summary_rows = [
        ["Generated On", payload["generated_on"]],
        ["Generated By", payload["generated_by"]],
    ]
    for office_key in payload["office_keys"]:
        office_data = payload["offices"][office_key]
        summary_rows.extend([
            [f"{office_data['label']} Manual Receipts Amount", office_data["manual_completed_total"]],
            [f"{office_data['label']} Imported Spreadsheet Receipt Amount", office_data["imported_total"]],
            [f"{office_data['label']} Imported Spreadsheet Rows Processed", office_data["imported_count"]],
            [f"{office_data['label']} Combined Total", office_data["combined_total"]],
            [f"{office_data['label']} Pending Manual Receipts", office_data["manual_pending_count"]],
            [f"{office_data['label']} Flagged Date Issues", office_data["date_quality_issue_count"]],
        ])
    if payload.get("combined"):
        summary_rows.extend([
            ["Combined Manual Receipts", payload["combined"]["manual_completed_total"]],
            ["Combined Imported Receipts", payload["combined"]["imported_total"]],
            ["Combined Financial Total", payload["combined"]["combined_total"]],
        ])
    _append_table(summary, "Financial Executive Summary", ["Metric", "Value"], summary_rows)

    for office_key in payload["office_keys"]:
        office_data = payload["offices"][office_key]
        overview = wb.create_sheet(f"{office_data['label']} Overview"[:31])
        _style_sheet(overview, GOV_GREEN if office_key == "nursing" else GOV_NAVY)
        _append_table(overview, f"{office_data['label']} Financial Summary", ["Metric", "Value"], _financial_summary_rows(office_data))
        _append_table(
            overview,
            f"{office_data['label']} Receipt Categories",
            ["Category", "Amount (PGK)"],
            [[row["label"], row["amount"]] for row in office_data["category_rows"]],
        )
        _append_table(
            overview,
            f"{office_data['label']} How To Read This Report",
            ["Guideline"],
            [[row] for row in office_data["guide_rows"]],
        )

        monthly_ws = wb.create_sheet(f"{office_data['label']} Monthly"[:31])
        _style_sheet(monthly_ws, GOV_GOLD if office_key == "nursing" else GOV_GREEN)
        month_start, month_end = _append_table(
            monthly_ws,
            f"{office_data['label']} Monthly Financial Tracking",
            [
                "Month",
                "Manual Amount",
                "Imported Amount",
                "Combined Total",
                "Manual Count",
                "Imported Count",
                "Renewals",
                "Full Registration",
                "New Applications",
                "Temporary",
                "Other",
            ],
            [
                [
                    row["period"],
                    row["manual_amount"],
                    row["imported_amount"],
                    row["total_amount"],
                    row["manual_count"],
                    row["imported_count"],
                    row["renewal_amount"],
                    row["full_registration_amount"],
                    row["new_application_amount"],
                    row["temporary_amount"],
                    row["other_amount"],
                ]
                for row in office_data["monthly_rows"]
            ],
        )
        _add_line_chart(monthly_ws, f"{office_data['label']} Monthly Total", month_start, month_end, [4], "M3")
        forecast_start, forecast_end = _append_table(
            monthly_ws,
            f"{office_data['label']} Forecast",
            ["Forecast Month", "Projected Amount", "Lower Range", "Upper Range"],
            [
                [row["period"], row["forecast_amount"], row["low_amount"], row["high_amount"]]
                for row in office_data["forecast_rows"]
            ],
        )
        _add_bar_chart(monthly_ws, f"{office_data['label']} Forecast Outlook", forecast_start, forecast_end, 2, "M22")

        yearly_ws = wb.create_sheet(f"{office_data['label']} Yearly"[:31])
        _style_sheet(yearly_ws, GOV_NAVY)
        year_start, year_end = _append_table(
            yearly_ws,
            f"{office_data['label']} Yearly Financial Tracking",
            [
                "Year",
                "Manual Amount",
                "Imported Amount",
                "Combined Total",
                "Manual Count",
                "Imported Count",
                "Renewals",
                "Full Registration",
                "New Applications",
                "Temporary",
                "Other",
            ],
            [
                [
                    row["period"],
                    row["manual_amount"],
                    row["imported_amount"],
                    row["total_amount"],
                    row["manual_count"],
                    row["imported_count"],
                    row["renewal_amount"],
                    row["full_registration_amount"],
                    row["new_application_amount"],
                    row["temporary_amount"],
                    row["other_amount"],
                ]
                for row in office_data["yearly_rows"]
            ],
        )
        _add_bar_chart(yearly_ws, f"{office_data['label']} Yearly Total", year_start, year_end, 4, "M3")

        recent_ws = wb.create_sheet(f"{office_data['label']} Transactions"[:31])
        _style_sheet(recent_ws, GOV_GREEN)
        _append_table(
            recent_ws,
            f"{office_data['label']} Recent Receipt Transactions",
            ["Source", "Reporting Date", "Source Date Issue", "Category", "Reference", "Amount (PGK)", "Status", "Detail"],
            [
                [
                    row["source"],
                    row["date"],
                    row["date_note"],
                    row["category"],
                    row["reference"],
                    row["amount"],
                    row["status"],
                    row["detail"],
                ]
                for row in office_data["recent_transactions"]
            ],
        )
        if office_data["date_quality_issues"]:
            _append_table(
                recent_ws,
                f"{office_data['label']} Flagged Date Issues",
                ["Source", "Original Source Date", "Reporting Date Used", "Reason", "Reference", "Amount (PGK)", "Detail", "Record ID"],
                [
                    [
                        row["source"],
                        row["source_date"],
                        row["reporting_date"],
                        row["reason"],
                        row["reference"],
                        row["amount"],
                        row["detail"],
                        row["record_id"],
                    ]
                    for row in office_data["date_quality_issues"]
                ],
            )

    output = BytesIO()
    wb.save(output)
    return output.getvalue()


def build_financial_forecast_pdf(office=None, generated_by=None):
    payload = build_financial_forecast_payload(office, generated_by=generated_by)
    output = BytesIO()
    doc = SimpleDocTemplate(output, pagesize=landscape(letter), leftMargin=24, rightMargin=24, topMargin=52, bottomMargin=28)
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle("FinanceTitle", parent=styles["Heading1"], fontSize=16, leading=19, textColor=colors.HexColor(f"#{GOV_NAVY}"), alignment=1, spaceAfter=4)
    subtitle_style = ParagraphStyle("FinanceSubtitle", parent=styles["Normal"], fontSize=9, leading=11, textColor=colors.HexColor(f"#{GOV_GOLD}"), alignment=1, spaceAfter=8)
    section_style = ParagraphStyle("FinanceSection", parent=styles["Heading2"], fontSize=11, leading=13, textColor=colors.HexColor(f"#{GOV_GREEN}"), spaceBefore=8)
    normal = ParagraphStyle("FinanceNormal", parent=styles["Normal"], fontSize=8, leading=10)

    story = [
        Paragraph("National Department of Health", title_style),
        Paragraph("Regulatory Bodies Financial Forecast and Receipt Tracking Report", subtitle_style),
        Paragraph(
            f"Generated on {payload['generated_on']} by {payload['generated_by']}. This report separates Nursing Council and Medical Board financial tracking using live receipts and spreadsheet-imported payment rows.",
            normal,
        ),
        Spacer(1, 8),
    ]

    summary_rows = []
    for office_key in payload["office_keys"]:
        office_data = payload["offices"][office_key]
        summary_rows.extend([
            [f"{office_data['label']} Manual Receipts", office_data["manual_completed_total"]],
            [f"{office_data['label']} Imported Spreadsheet Receipt Amount", office_data["imported_total"]],
            [f"{office_data['label']} Imported Spreadsheet Rows Processed", office_data["imported_count"]],
            [f"{office_data['label']} Combined Total", office_data["combined_total"]],
            [f"{office_data['label']} Pending Manual Receipts", office_data["manual_pending_count"]],
            [f"{office_data['label']} Flagged Date Issues", office_data["date_quality_issue_count"]],
        ])
    story.append(_pdf_table(["Metric", "Value"], summary_rows, doc.width))

    for office_key in payload["office_keys"]:
        office_data = payload["offices"][office_key]
        story.extend([
            Spacer(1, 10),
            Paragraph(office_data["label"], section_style),
            _pdf_table(["Metric", "Value"], _financial_summary_rows(office_data), doc.width * 0.52),
            Spacer(1, 6),
            _pdf_table(["Category", "Amount (PGK)"], [[row["label"], row["amount"]] for row in office_data["category_rows"]], doc.width * 0.52),
            Spacer(1, 6),
            _pdf_table(
                ["Month", "Manual", "Imported", "Combined", "Renewals", "Full Registration", "New Applications"],
                [
                    [
                        row["period"],
                        row["manual_amount"],
                        row["imported_amount"],
                        row["total_amount"],
                        row["renewal_amount"],
                        row["full_registration_amount"],
                        row["new_application_amount"],
                    ]
                    for row in office_data["monthly_rows"][-12:]
                ],
                doc.width,
            ),
            Spacer(1, 6),
            _pdf_table(
                ["Forecast Month", "Projected Amount", "Lower Range", "Upper Range"],
                [
                    [row["period"], row["forecast_amount"], row["low_amount"], row["high_amount"]]
                    for row in office_data["forecast_rows"]
                ],
                doc.width * 0.58,
            ),
            Spacer(1, 6),
            _pdf_table(
                ["Guideline"],
                [[row] for row in office_data["guide_rows"]],
                doc.width,
            ),
        ])
        if office_data["date_quality_issues"]:
            story.extend([
                Spacer(1, 6),
                Paragraph("Flagged source-date issues", section_style),
                _pdf_table(
                    ["Source", "Original Date", "Reporting Date", "Reason", "Reference", "Amount"],
                    [
                        [
                            row["source"],
                            row["source_date"],
                            row["reporting_date"],
                            row["reason"],
                            row["reference"],
                            row["amount"],
                        ]
                        for row in office_data["date_quality_issues"][:12]
                    ],
                    doc.width,
                ),
            ])

    doc.build(story, onFirstPage=_draw_government_pdf_frame, onLaterPages=_draw_government_pdf_frame)
    return output.getvalue()


def build_financial_forecast_docx(office=None, generated_by=None):
    payload = build_financial_forecast_payload(office, generated_by=generated_by)
    doc = WordDocument()
    section = doc.sections[0]
    section.orientation = WD_ORIENTATION.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.45)
    section.right_margin = Inches(0.45)

    logo_path = _ndoh_logo_path()
    if logo_path:
        logo_paragraph = doc.add_paragraph()
        logo_paragraph.alignment = 1
        logo_run = logo_paragraph.add_run()
        try:
            logo_run.add_picture(str(logo_path), width=Inches(0.7))
        except Exception:
            pass

    heading = doc.add_paragraph()
    run = heading.add_run("NATIONAL DEPARTMENT OF HEALTH")
    run.bold = True
    run.font.size = Pt(16)
    heading.alignment = 1

    subheading = doc.add_paragraph()
    run = subheading.add_run("Regulatory Bodies Financial Forecast and Receipt Tracking Report")
    run.bold = True
    run.font.size = Pt(12)
    subheading.alignment = 1

    meta = doc.add_paragraph(f"Generated on {payload['generated_on']} by {payload['generated_by']}. This report separates Nursing Council and Medical Board receipt tracking using manual receipts and spreadsheet-imported payment records.")
    meta.alignment = 1

    doc.add_paragraph("")
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Metric"
    hdr[1].text = "Value"
    for office_key in payload["office_keys"]:
        office_data = payload["offices"][office_key]
        for label, value in _financial_summary_rows(office_data):
            row = table.add_row().cells
            row[0].text = f"{office_data['label']} - {label}"
            row[1].text = str(value)

    for office_key in payload["office_keys"]:
        office_data = payload["offices"][office_key]
        doc.add_page_break()
        title = doc.add_paragraph()
        title_run = title.add_run(office_data["label"])
        title_run.bold = True
        title_run.font.size = Pt(14)

        guide_heading = doc.add_paragraph()
        guide_heading.add_run("How to understand this financial report").bold = True
        for guide_row in office_data["guide_rows"]:
            doc.add_paragraph(guide_row, style="List Bullet")

        category_table = doc.add_table(rows=1, cols=2)
        category_table.style = "Table Grid"
        category_table.rows[0].cells[0].text = "Receipt Category"
        category_table.rows[0].cells[1].text = "Amount (PGK)"
        for row_data in office_data["category_rows"]:
            row = category_table.add_row().cells
            row[0].text = row_data["label"]
            row[1].text = str(row_data["amount"])

        doc.add_paragraph("")
        monthly_table = doc.add_table(rows=1, cols=7)
        monthly_table.style = "Table Grid"
        headers = ["Month", "Manual", "Imported", "Combined", "Renewals", "Full Registration", "New Applications"]
        for idx, header in enumerate(headers):
            monthly_table.rows[0].cells[idx].text = header
        for row_data in office_data["monthly_rows"][-12:]:
            row = monthly_table.add_row().cells
            row[0].text = str(row_data["period"])
            row[1].text = str(row_data["manual_amount"])
            row[2].text = str(row_data["imported_amount"])
            row[3].text = str(row_data["total_amount"])
            row[4].text = str(row_data["renewal_amount"])
            row[5].text = str(row_data["full_registration_amount"])
            row[6].text = str(row_data["new_application_amount"])

        doc.add_paragraph("")
        forecast_table = doc.add_table(rows=1, cols=4)
        forecast_table.style = "Table Grid"
        headers = ["Forecast Month", "Projected Amount", "Lower Range", "Upper Range"]
        for idx, header in enumerate(headers):
            forecast_table.rows[0].cells[idx].text = header
        for row_data in office_data["forecast_rows"]:
            row = forecast_table.add_row().cells
            row[0].text = row_data["period"]
            row[1].text = str(row_data["forecast_amount"])
            row[2].text = str(row_data["low_amount"])
            row[3].text = str(row_data["high_amount"])

        if office_data["date_quality_issues"]:
            doc.add_paragraph("")
            issue_heading = doc.add_paragraph()
            issue_heading.add_run("Flagged source-date issues").bold = True
            issue_table = doc.add_table(rows=1, cols=6)
            issue_table.style = "Table Grid"
            issue_headers = ["Source", "Original Date", "Reporting Date", "Reason", "Reference", "Amount"]
            for idx, header in enumerate(issue_headers):
                issue_table.rows[0].cells[idx].text = header
            for issue in office_data["date_quality_issues"][:20]:
                row = issue_table.add_row().cells
                row[0].text = str(issue["source"])
                row[1].text = str(issue["source_date"])
                row[2].text = str(issue["reporting_date"])
                row[3].text = str(issue["reason"])
                row[4].text = str(issue["reference"])
                row[5].text = str(issue["amount"])

    output = BytesIO()
    doc.save(output)
    return output.getvalue()
