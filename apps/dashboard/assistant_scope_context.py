from pathlib import Path

from django.core.cache import cache

from apps.dashboard.nursing_analytics import active_nursing_analytics_snapshot, metric_payload
from apps.dashboard.nursing_catherine_breakdown import catherine_breakdown_overlay_payload


NURSING_CURRENT_SOURCE_DIR = (
    Path.home()
    / "Documents"
    / "ProjectApps"
    / "databasedocuments"
    / "spreadsheets"
    / "Nursing_Council_Cleansed_data_current"
)
NURSING_GUIDELINE_DOC = NURSING_CURRENT_SOURCE_DIR / "PNG_Nursing_Council_Excel_Template_Guideline_and_Update.docx"

NURSING_OFFICE_TERMS = (
    "nursing",
    "nurse",
    "midwife",
    "midwifery",
    "graduand",
    "nurse aide",
    "provisional",
    "full licence",
    "full license",
    "atp",
    "authority to practice",
    "nc1",
    "nc2",
    "nc3",
    "nc4",
    "nc5",
    "nc6",
    "nc7",
    "nc8",
    "nc9",
)
MEDICAL_OFFICE_TERMS = (
    "medical board",
    "doctor",
    "medical practitioner",
    "specialist",
    "chw",
    "community health worker",
    "medical staff",
    "md1",
    "md2",
    "mbsp",
    "mbrn",
    "mbac",
    "mbpf",
    "mbtc",
)


def scope_policy_context(scope):
    if scope == "nursing":
        allowed = "Nursing Council records, forms, pathways, licence data, ATP renewals, cadre breakdowns, and Nursing Council reports."
        restricted = "Medical Board doctor, specialist, CHW, and Medical Board facility records."
    elif scope == "medical":
        allowed = "Medical Board doctor, specialist, CHW, facility, medical application, and Medical Board report records."
        restricted = "Nursing Council nurse, midwife, nurse aide, graduand, ATP, and Nursing Council analytics records."
    elif scope == "all":
        allowed = "Both Nursing Council and Medical Board operational summaries, with office separation preserved in every answer."
        restricted = "Private practitioner details outside the signed-in user's authorised role and any final registrar decision made by AI."
    else:
        allowed = "Public-safe platform guidance only."
        restricted = "Private practitioner, staff, import, and registrar-only records."

    return {
        "scope": scope or "public",
        "allowed": allowed,
        "restricted": restricted,
        "rules": [
            "Answer from the signed-in user's regulatory office scope first.",
            "If a question belongs to the other office, explain the correct office instead of mixing data.",
            "Never approve, reject, merge, mark deceased, or update records from an AI response.",
            "Use Nursing Council ATP/lapsed signals only as registrar review candidates.",
        ],
    }


def detect_cross_scope_question(scope, question):
    text = " ".join(str(question or "").lower().split())
    has_nursing = any(term in text for term in NURSING_OFFICE_TERMS)
    has_medical = any(term in text for term in MEDICAL_OFFICE_TERMS)
    asks_boundary_policy = any(
        term in text
        for term in ("scope", "privacy", "access", "separate", "separation", "boundary", "role", "roles", "who can view")
    )
    if asks_boundary_policy and has_nursing and has_medical:
        return {"detected": False}
    if scope == "nursing" and has_medical:
        return {
            "detected": True,
            "target_scope": "medical",
            "message": "That question belongs to the Medical Board scope. Nursing Council staff should not use Nursing Council analytics to answer doctor, specialist, CHW, or Medical Board facility questions.",
        }
    if scope == "medical" and has_nursing:
        return {
            "detected": True,
            "target_scope": "nursing",
            "message": "That question belongs to the Nursing Council scope. Medical Board staff should not use Medical Board data to answer nurse, midwife, nurse aide, graduand, ATP, or Nursing Council pathway questions.",
        }
    return {"detected": False}


def _workflow_definitions():
    try:
        from apps.workforce.services.nursing_council_workflows import FORM_DEFINITIONS, PATHWAY_DEFINITIONS
    except Exception:
        return [], []
    return FORM_DEFINITIONS, PATHWAY_DEFINITIONS


def nursing_pathway_context():
    form_definitions, pathway_definitions = _workflow_definitions()
    pathways = []
    for row in sorted(pathway_definitions, key=lambda item: item.get("sort_order", 999)):
        pathways.append({
            "pathway_code": row.get("pathway_code", ""),
            "pathway_name": row.get("pathway_name", ""),
            "primary_form_code": row.get("primary_form_code", ""),
            "checklist_code": row.get("checklist_code", ""),
            "applicant_type": row.get("applicant_type", ""),
            "creates_licence_type": row.get("creates_licence_type", ""),
            "requires_payment": bool(row.get("requires_payment")),
            "requires_registrar_approval": bool(row.get("requires_registrar_approval")),
        })

    forms = [
        {
            "form_code": form_code,
            "form_name": form_name,
            "pathway_code": pathway_code,
        }
        for form_code, form_name, pathway_code in form_definitions
    ]
    return {
        "source": "apps.workforce.services.nursing_council_workflows",
        "pathways": pathways,
        "forms": forms,
        "plain_language": [
            "NC1 is the main provisional licence entry path for PNG graduate nurses and overseas provisional applicants.",
            "NC2 supports full registration, including provisional-to-full and midwifery full-registration pathways.",
            "NC3 is the renewal pathway and aligns with Authority to Practice/practising renewal tracking.",
            "NC8 and NC9 support temporary overseas licence applications and the temporary overseas checklist.",
            "ATP records are renewal/practising evidence; they are not the same thing as initial full-registration application rows.",
        ],
    }


def _docx_dataflow_summary():
    cache_key = f"nursing-guideline-doc:{NURSING_GUIDELINE_DOC}:{NURSING_GUIDELINE_DOC.stat().st_mtime if NURSING_GUIDELINE_DOC.exists() else 'missing'}"
    cached = cache.get(cache_key)
    if cached is not None:
        return cached

    summary = {
        "document": NURSING_GUIDELINE_DOC.name,
        "available": NURSING_GUIDELINE_DOC.exists(),
        "highlights": [],
        "workbook_roles": [],
    }
    if not NURSING_GUIDELINE_DOC.exists():
        cache.set(cache_key, summary, 300)
        return summary

    try:
        from docx import Document
        document = Document(NURSING_GUIDELINE_DOC)
    except Exception as exc:
        summary["highlights"] = [f"Could not read guideline document: {exc}"]
        cache.set(cache_key, summary, 300)
        return summary

    for paragraph in document.paragraphs:
        text = " ".join(paragraph.text.split())
        lowered = text.lower()
        if not text:
            continue
        if any(token in lowered for token in ("cleaned", "review", "approval", "dashboard", "data quality")):
            summary["highlights"].append(text[:220])
        if len(summary["highlights"]) >= 6:
            break

    for table in document.tables:
        if len(table.rows) < 2:
            continue
        headers = [cell.text.strip().lower() for cell in table.rows[0].cells]
        if "file name" not in " ".join(headers) and "input / output" not in " ".join(headers):
            continue
        for row in table.rows[1:8]:
            cells = [" ".join(cell.text.split()) for cell in row.cells]
            if len(cells) >= 3 and any(cells):
                summary["workbook_roles"].append({
                    "input_output": cells[0][:100],
                    "file_name": cells[1][:120],
                    "purpose": cells[2][:220],
                })
        if summary["workbook_roles"]:
            break

    cache.set(cache_key, summary, 300)
    return summary


def nursing_cadre_dataflow_context(limit=12):
    snapshot = active_nursing_analytics_snapshot()
    payload = metric_payload(snapshot)
    overlay = catherine_breakdown_overlay_payload(snapshot)
    kpis = payload.get("kpis", {})
    live_stats = payload.get("live_statistics", {})
    cadre_rows = overlay.get("cadre_rows") or []
    if not cadre_rows:
        chart = (payload.get("charts") or {}).get("cadre") or {}
        labels = chart.get("labels", [])
        cadre_rows = [
            {
                "cadre": label,
                "provisional_clean": (chart.get("provisional") or [])[index] if index < len(chart.get("provisional") or []) else 0,
                "full_professional_clean": (chart.get("full_licence") or [])[index] if index < len(chart.get("full_licence") or []) else 0,
                "authority_to_practice_count": (chart.get("authority_to_practice") or [])[index] if index < len(chart.get("authority_to_practice") or []) else 0,
                "total_clean": 0,
                "profession_group": "",
                "classification_note": "",
            }
            for index, label in enumerate(labels)
        ]

    top_rows = sorted(cadre_rows, key=lambda row: int(row.get("total_clean") or row.get("grand_total") or 0), reverse=True)[:limit]
    source_files = overlay.get("source_files") or {}
    doc_summary = _docx_dataflow_summary()
    return {
        "source": "Active Nursing Council analytics snapshot and current Nursing Council cleansed workbook guide.",
        "snapshot": {
            "id": snapshot.pk if snapshot else None,
            "source_file_name": snapshot.source_file_name if snapshot else "",
            "source_file_hash": snapshot.source_file_hash if snapshot else "",
            "workbook_generated_on": snapshot.workbook_generated_on.isoformat() if snapshot and snapshot.workbook_generated_on else "",
        },
        "source_files": {
            "integrated_dashboard": "PNG_Nursing_Council_Integrated_Dashboard_Model.xlsx",
            "cleaned_licence": (source_files.get("cleaned_licence") or {}).get("file_name", "PNG_Nursing_Council_Cleaned_Licence_Breakdown.xlsx"),
            "cadre_breakdown": (source_files.get("cadre_breakdown") or {}).get("file_name", "PNG_Nursing_Council_Cadre_Breakdown.xlsx"),
            "guideline": NURSING_GUIDELINE_DOC.name,
        },
        "totals": {
            "total_lifecycle_records": int(kpis.get("total_lifecycle_records") or 0),
            "clean_provisional_records": int(kpis.get("clean_provisional_records") or 0),
            "clean_full_licence_records": int(kpis.get("clean_full_licence_records") or 0),
            "clean_atp_records": int(kpis.get("clean_atp_records") or 0),
            "current_atp_year": live_stats.get("atp_current_year"),
            "current_atp_people": live_stats.get("atp_current_person_total"),
            "mapped_clean_rows": overlay.get("mapped_clean_rows", 0),
            "unclassified_clean_rows": overlay.get("unclassified_clean_rows", 0),
        },
        "cadre_rows": top_rows,
        "dataflow_steps": [
            "Legacy provisional and full-licence data is cleaned into the licence breakdown workbook.",
            "Qualification text is mapped through the cadre breakdown workbook into cadre groups such as Nursing Graduand, Registered Nurse, Midwifery, and review/unclassified groups.",
            "The integrated dashboard model consolidates lifecycle facts, practitioner index rows, stage/year metrics, cadre/stage metrics, facility metrics, institution metrics, and data-quality metrics.",
            "The Django platform imports the integrated dashboard model as an analytics snapshot, not as automatic legal registry approval.",
            "Catherine's cleaned licence and cadre workbooks are attached as a verification overlay against the active snapshot and should not be double-counted.",
        ],
        "guideline_document": doc_summary,
    }


def nursing_cadre_answer_payload(context):
    cadre_context = context.get("nursing_cadre_context") or {}
    totals = cadre_context.get("totals") or {}
    rows = cadre_context.get("cadre_rows") or []
    row_lines = []
    for row in rows[:6]:
        row_lines.append(
            f"{row.get('cadre')}: provisional {row.get('provisional_clean', row.get('provisional_licence_count', 0))}, "
            f"full {row.get('full_professional_clean', row.get('full_licence_count', 0))}, "
            f"total {row.get('total_clean', row.get('grand_total', 0))}"
        )
    return {
        "answer": (
            "The Nursing Council cadre breakdown links the licence pathway to the cleaned analytics model: "
            f"{totals.get('clean_provisional_records', 0)} clean provisional records, "
            f"{totals.get('clean_full_licence_records', 0)} clean full-licence records, and "
            f"{totals.get('clean_atp_records', 0)} ATP/practising records."
        ),
        "bullets": row_lines,
    }
