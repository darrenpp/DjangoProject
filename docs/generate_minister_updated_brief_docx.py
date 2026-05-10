from pathlib import Path
import csv
import os
import shutil
import sys

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

PROJECT_NAME = "The National Department Of Health Regulatory Bodies Nursing Council & The Medical Board Online Workforce System"
SUBJECT = f"Updated Brief for {PROJECT_NAME}."
BASE_DIR = Path(__file__).resolve().parent.parent
if str(BASE_DIR) not in sys.path:
    sys.path.insert(0, str(BASE_DIR))
os.environ.setdefault("DJANGO_SETTINGS_MODULE", "ndoh_workforce_registry.settings")

import django

django.setup()

from apps.dashboard.reports import build_monthly_analytics_payload
from apps.workforce.models import EmploymentRecord, HealthStudent, Midwife, NurseAide, NursingProfessional

DOCS_DIR = BASE_DIR / "docs"
NOTEBOOKS_DIR = BASE_DIR / "notebooks"
ASSETS_DIR = DOCS_DIR / "system_brief_assets"
LOCAL_OUTPUT_PATH = DOCS_DIR / "NDOH_Regulatory_Bodies_Online_Workforce_System_Brief_Minister_Updated.docx"
EXTERNAL_OUTPUT_PATH = Path(r"c:\Users\timhi\OneDrive\Desktop\ParotOs\NDOH_Database\Briefs\march_briefs_2026\NDOH_Regulatory_Bodies_Online_Workforce_System_Brief_Minister_Updated.docx")


def read_summary(path: Path):
    result = {}
    if not path.exists():
        return result
    for line in path.read_text(encoding="utf-8").splitlines():
        if ":" not in line:
            continue
        key, value = line.split(":", 1)
        result[key.strip()] = value.strip()
    return result


def read_issue_counts(path: Path, issue_column: str):
    counts = {}
    if not path.exists():
        return counts
    with path.open(newline="", encoding="utf-8") as handle:
        reader = csv.DictReader(handle)
        for row in reader:
            issue = (row.get(issue_column) or "").strip()
            if not issue:
                continue
            counts[issue] = counts.get(issue, 0) + 1
    return counts


def set_default_font(document: Document):
    style = document.styles["Normal"]
    style.font.name = "Times New Roman"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    style.font.size = Pt(11)

    for style_name in ["Title", "Heading 1", "Heading 2", "Heading 3"]:
        if style_name in document.styles:
            style = document.styles[style_name]
            style.font.name = "Times New Roman"
            style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")


def add_bottom_border(paragraph):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "18")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "000000")
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def add_bullet(document: Document, text: str):
    paragraph = document.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.add_run(text)


def add_number(document: Document, text: str):
    paragraph = document.add_paragraph(style="List Number")
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.add_run(text)


def add_table(document: Document, headers, rows):
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.rows[0].cells[0].paragraphs[0].runs
    for cell, text in zip(table.rows[0].cells, headers):
        cell.text = str(text)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        for run in cell.paragraphs[0].runs:
            run.bold = True
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cells[index].text = "" if value is None else str(value)
            cells[index].vertical_alignment = WD_ALIGN_VERTICAL.TOP
    return table


def add_large_screenshot(document: Document, title: str, image_name: str, caption: str, explanation_lines):
    document.add_heading(title, level=2)
    image_path = ASSETS_DIR / image_name
    if image_path.exists():
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.add_run().add_picture(str(image_path), width=Inches(5.45))
        caption_paragraph = document.add_paragraph(caption)
        caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if caption_paragraph.runs:
            caption_paragraph.runs[0].italic = True
            caption_paragraph.runs[0].font.size = Pt(9)
    details = document.add_paragraph()
    details.add_run("What the Minister should note on this screen").bold = True
    for line in explanation_lines:
        add_bullet(document, line)


def build_live_summary():
    payload = build_monthly_analytics_payload("nursing")["nursing"]
    latest_batch = payload["latest_batch_row"] or {}
    source_sheets = payload["source_sheet_rows"]
    recent_sheet_text = ", ".join(
        f"{row['sheet_name'].strip()} ({row['imported_rows']} rows)"
        for row in source_sheets[:5]
    )

    employed_count = (
        EmploymentRecord.objects.exclude(employment_status="")
        .exclude(employment_status__isnull=True)
        .exclude(employment_status__in=["unemployed", "studying"])
        .count()
    )
    unemployed_count = EmploymentRecord.objects.filter(employment_status="unemployed").count()

    headline_rows = [
        ["Total Registered Nurses", NursingProfessional.objects.count(), "Current live count of registered nurses in the electronic registry."],
        ["Total Midwives", Midwife.objects.count(), "Current live count of midwives in the electronic registry."],
        ["Total Nurse Aides", NurseAide.objects.count(), "Current live count of nurse aides in the electronic registry."],
        ["Total Employed", employed_count, "Employment records captured electronically as employed. Current value remains low because the employment module is not yet populated."],
        ["Total Unemployed", unemployed_count, "Employment records captured electronically as unemployed. Current value remains low because the employment module is not yet populated."],
    ]

    registry_rows = [
        [row["label"], row["current_count"], row["active_count"], row["latest_updated"]]
        for row in payload["current_registry_rows"]
    ]

    source_rows = [
        ["Latest source file", latest_batch.get("source_file_name", "Not captured")],
        ["Latest completed import", latest_batch.get("completed_at", "Not captured")],
        ["Source kind", latest_batch.get("source_kind", "Not captured")],
        ["Imported rows in latest batch", latest_batch.get("processed_rows", "Not captured")],
        ["Latest source sheets", recent_sheet_text or "Not captured"],
        ["Current live registry people", payload["current_registry_total"]],
        ["Imported record rows in analytics", payload["total_records"]],
    ]

    application_rows = [
        [row["status"].title(), row["count"], "Current application status total for Nursing Council forms."]
        for row in payload["application_rows"]
    ]

    record_mix_rows = [
        [row["record_type"].replace("_", " ").title(), row["count"], "Imported historical record count in the Nursing Council analytics dataset."]
        for row in payload["record_type_rows"]
    ]

    return {
        "headline_rows": headline_rows,
        "registry_rows": registry_rows,
        "source_rows": source_rows,
        "application_rows": application_rows,
        "record_mix_rows": record_mix_rows,
    }


def build_document():
    doc = Document()
    set_default_font(doc)

    section = doc.sections[0]
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)

    crest = DOCS_DIR.parent / "static" / "img" / "NDOH_LOGO.png"
    if crest.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(crest), width=Inches(1.8))

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("PAPUA NEW GUINEA NURSING COUNCIL\nOFFICE OF THE REGISTRAR")
    run.bold = True
    run.font.size = Pt(15)

    border = doc.add_paragraph()
    border.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_bottom_border(border)

    doc.add_paragraph()
    salutation = doc.add_paragraph()
    salutation.add_run("Dear Hon. Minister for Health,").bold = True

    subject = doc.add_paragraph()
    subject_run = subject.add_run(f"Subject: {SUBJECT}")
    subject_run.bold = True
    subject_run.underline = True

    intro = [
        f"This updated brief provides a plain-language overview of {PROJECT_NAME}. It has been rewritten for executive reading and uses the current live system interface images, larger visual layouts, and the latest available live registry and analytics statistics.",
        "The document covers the full scope of the platform: public entry screens, Nursing Council functions, Medical Board functions, self-service user portals, registrar workflows, reporting, privacy controls, and the quality requirements for building a clean national workforce register.",
    ]
    for text in intro:
        paragraph = doc.add_paragraph(text)
        paragraph.paragraph_format.space_after = Pt(6)

    doc.add_heading("1. System Name and Purpose", level=1)
    doc.add_paragraph(
        f"The official project name used in this brief is: {PROJECT_NAME}. In practical terms, it is an online workforce management, registration, and reporting platform for the Nursing Council and the Medical Board under the National Department of Health."
    )
    doc.add_paragraph(
        "The system is intended to support application intake, practitioner registration, licence monitoring, qualification storage, document review, data-cleansing, and management reporting."
    )

    doc.add_heading("2. Executive Summary in Plain Language", level=1)
    for item in [
        "The system is both an operational processing tool and a long-term regulatory record system.",
        "It separates Nursing Council and Medical Board information so that each regulator stays inside its own approved data space.",
        "It supports public form intake, registrar review, applicant self-service, document storage, historical imports, dashboards, and reporting.",
        "It is already useful for daily work, but the quality of incoming paper records and legacy spreadsheets still determines the quality of the final electronic register.",
    ]:
        add_bullet(doc, item)

    doc.add_heading("3. Summary of current updated data", level=1)
    doc.add_paragraph(
        "This section summarises the latest live Nursing Council statistics now available from the current electronic registry and the Nursing Council monthly analytics report. These are not comparison figures. They are the current live figures available in the platform at the time this brief was prepared."
    )
    live = build_live_summary()
    add_table(doc, ["Live statistic", "Current total", "Meaning"], live["headline_rows"])
    doc.add_paragraph()
    note = doc.add_paragraph()
    note.add_run("Important note on employed and unemployed totals: ").bold = True
    note.add_run(
        "the employment module exists in the system, but the live Employment Record table is not yet populated. The electronic totals for employed and unemployed therefore remain at zero until staff update those records consistently."
    )

    doc.add_paragraph()
    doc.add_paragraph().add_run("Current live registry snapshot").bold = True
    add_table(doc, ["Registry category", "Current live count", "Active count", "Latest update in system"], live["registry_rows"])

    doc.add_paragraph()
    doc.add_paragraph().add_run("Latest source update behind the live nursing analytics").bold = True
    add_table(doc, ["Metric", "Value"], live["source_rows"])

    doc.add_paragraph()
    doc.add_paragraph().add_run("Current Nursing Council application status totals").bold = True
    add_table(doc, ["Application status", "Current total", "Meaning"], live["application_rows"])

    doc.add_paragraph()
    doc.add_paragraph().add_run("Current Nursing Council record activity mix from the monthly analytics dataset").bold = True
    add_table(doc, ["Record activity", "Current total", "Meaning"], live["record_mix_rows"])

    doc.add_heading("4. Main interfaces with current system screenshots", level=1)
    doc.add_paragraph(
        "The screenshots below are included to help the Minister see the actual user interfaces currently used inside the platform. They have been enlarged for readability and paired with a short explanation of what each screen is used for."
    )

    add_large_screenshot(
        doc,
        "4.1 Public home page",
        "public_home.png",
        "Current public entry screen",
        [
            "This is the front entry point into the platform.",
            "It directs different users to the correct application or login pathway.",
            "It is important because it reduces confusion and helps separate Nursing Council and Medical Board user journeys from the beginning.",
        ],
    )
    add_large_screenshot(
        doc,
        "4.2 Overall dashboard",
        "overall_dashboard_admin.png",
        "Current overall leadership dashboard",
        [
            "This screen gives senior users a combined high-level view of activity across the platform.",
            "It is useful for leadership oversight, but role-based restrictions still control what each user can actually access.",
            "The dashboard helps managers understand volume, workflow, and overall system use.",
        ],
    )
    add_large_screenshot(
        doc,
        "4.3 Nursing Council dashboard",
        "nursing_dashboard.png",
        "Current Nursing Council operational dashboard",
        [
            "This is the main Nursing Council operational workspace.",
            "It focuses on nursing, midwifery, nurse aide, and graduand information.",
            "It supports reporting, searches, form access, and current nursing statistics.",
        ],
    )
    add_large_screenshot(
        doc,
        "4.4 Medical Board dashboard",
        "medical_dashboard.png",
        "Current Medical Board operational dashboard",
        [
            "This is the Medical Board equivalent of the Nursing Council dashboard.",
            "It supports the Medical Board's own operational work and its own professional groups.",
            "It remains separate from the Nursing Council data environment.",
        ],
    )
    add_large_screenshot(
        doc,
        "4.5 Nursing forms portal",
        "nursing_forms.png",
        "Current Nursing Council forms portal",
        [
            "This screen directs applicants and staff to the correct Nursing Council forms.",
            "It helps reduce form errors and ensures the right workflow starts from the beginning.",
            "It also supports standardisation of required documents and fee pathways.",
        ],
    )
    add_large_screenshot(
        doc,
        "4.6 Medical Board forms portal",
        "medical_forms.png",
        "Current Medical Board forms portal",
        [
            "This screen serves the Medical Board form pathways.",
            "It keeps Medical Board forms separate from Nursing Council forms.",
            "This separation supports regulatory clarity and privacy control.",
        ],
    )
    add_large_screenshot(
        doc,
        "4.7 Professional self-service portal",
        "nurse_portal.png",
        "Current nurse self-service portal",
        [
            "This is an example of a personal portal used by an individual professional.",
            "It is intended for personal record viewing, application tracking, and limited self-service actions.",
            "Individual users should only view their own information through this pathway.",
        ],
    )
    add_large_screenshot(
        doc,
        "4.8 Individual professional record",
        "professional_detail.png",
        "Current professional record detail screen",
        [
            "This is one of the most important operational screens in the system.",
            "It brings together personal details, qualifications, documents, and related professional records in one place.",
            "It supports registrar review, verification, and evidence-based decision-making.",
        ],
    )
    add_large_screenshot(
        doc,
        "4.9 Graduand and student pathway",
        "graduand_portal.png",
        "Current graduand portal",
        [
            "This screen supports future practitioners moving from training into registration pathways.",
            "It helps track graduand progress, required actions, and related application steps.",
            "It is useful for workforce pipeline visibility and early regulatory tracking.",
        ],
    )
    add_large_screenshot(
        doc,
        "4.10 Workforce flow and planning view",
        "workforce_flow.png",
        "Current workforce flow dashboard",
        [
            "This interface supports management-level viewing of workforce movement and pipeline trends.",
            "It helps explain the flow from training, registration, and practice-related activities into the live registry picture.",
            "It supports planning and strategic reporting rather than simple data entry.",
        ],
    )
    add_large_screenshot(
        doc,
        "4.11 Administrative operations view",
        "admin_dashboard.png",
        "Current administrative oversight dashboard",
        [
            "This screen supports administrative and operational oversight tasks.",
            "It is useful for monitoring submissions, workflow, and processing support functions.",
            "It complements the regulator-specific dashboards with broader operational visibility.",
        ],
    )

    doc.add_heading("5. Roles, accessibilities, and privacy", level=1)
    doc.add_paragraph(
        "The platform is role-based. This means different users see different data depending on who they are, what office they belong to, and what level of permission they have been given."
    )
    add_table(
        doc,
        ["Role", "Main access", "What the role can view", "Privacy position"],
        [
            ["Admin", "Top-level oversight", "Can oversee system operations and technical management functions.", "This role is supervisory and must be tightly controlled."],
            ["Nursing Council registrar / staff", "Nursing Council workspace", "Can work on nursing, midwifery, nurse aide, graduand, forms, and related Nursing Council records.", "Should remain inside the Nursing Council data space."],
            ["Medical Board registrar / staff", "Medical Board workspace", "Can work on Medical Board records, forms, doctor records, and CHW records.", "Should remain inside the Medical Board data space."],
            ["Nurse / Nurse Aide / Doctor / CHW", "Personal portal", "Can view their own personal records, applications, and personal workflow items.", "Should not view another person's record."],
            ["Graduand", "Graduand portal", "Can view their own pathway, status, and related records.", "Should only see their own file."],
            ["Viewer / Reviewer", "Limited review access", "Can be given controlled access to support a specific workflow.", "Should stay within assigned tasks only."],
        ],
    )

    doc.add_paragraph()
    privacy_note = doc.add_paragraph()
    privacy_note.add_run("Privacy position: ").bold = True
    privacy_note.add_run(
        "The platform is intended to keep Nursing Council and Medical Board operational data separate. This supports confidentiality, proper regulatory handling, and professional governance."
    )

    doc.add_heading("6. What the system can do", level=1)
    for item in [
        "Receive online applications through separated Nursing Council and Medical Board form pathways.",
        "Store practitioner profiles for nurses, midwives, nurse aides, doctors, CHWs, and graduands.",
        "Store qualifications, institutions, completion years, and related education records.",
        "Store supporting files such as certificates, transcripts, ID documents, receipts, and photographs.",
        "Track application progress and review status.",
        "Support receipt submission and payment evidence capture.",
        "Support dashboards, exports, registry search, and management reports.",
        "Import historical workbook records into structured electronic records.",
        "Flag missing information and support data-quality review work.",
        "Support workflow tracking and registry decision support.",
    ]:
        add_bullet(doc, item)

    doc.add_heading("7. What the system is able to capture", level=1)
    add_table(
        doc,
        ["Data area", "Examples of information captured"],
        [
            ["Identity and contacts", "Names, registration numbers, email, phone, gender, date of birth, province."],
            ["Professional classification", "Nurse, midwife, nurse aide, doctor, CHW, graduand, cadre, applicant type."],
            ["Licence information", "Issue date, expiry date, registration pathway, provisional or full licence status."],
            ["Education and qualifications", "Institution, qualification title, programme completed, completion year, country."],
            ["Supporting records", "Certificates, transcripts, ID documents, passport-style images, uploaded files, receipt evidence."],
            ["Applications", "Form code, title, status, submission date, pathway details."],
            ["Historical registry records", "Imported workbook records, licence rows, workforce listing records, practice-related history."],
            ["Data quality tracking", "Missing fields, issue categories, review status, registrar follow-up items."],
        ],
    )

    doc.add_heading("8. Data quality findings and current challenges", level=1)
    full_summary = read_summary(NOTEBOOKS_DIR / "full_registrations_summary.txt")
    provisional_summary = read_summary(NOTEBOOKS_DIR / "provisional_graduands_summary.txt")
    full_issues = read_issue_counts(NOTEBOOKS_DIR / "full_registrations_quality_issues.csv", "issue_detail")
    prov_issues = read_issue_counts(NOTEBOOKS_DIR / "provisional_graduands_quality_issues.csv", "issues")

    doc.add_paragraph(
        "The system includes data-cleansing and review logic. It does not simply copy historical spreadsheets into the live registry without checks. Instead, it attempts to normalise records and flag what still needs review."
    )
    for metric in [
        f"Full-registration records cleaned: {full_summary.get('Cleaned records', 'N/A')}",
        f"Full-registration issue rows flagged: {full_summary.get('Quality issue rows', 'N/A')}",
        f"Provisional and graduand records cleaned: {provisional_summary.get('Cleaned records', 'N/A')}",
        f"Provisional and graduand issue rows flagged: {provisional_summary.get('Quality issue records', 'N/A')}",
    ]:
        add_bullet(doc, metric)

    doc.add_paragraph("The most common data challenges still being found include:")
    for item in [
        f"Invalid practitioner number formats ({full_issues.get('Invalid practitioner number format not imported', 0)} flagged rows in the full-registration issue file).",
        f"Duplicate practitioner numbers ({full_issues.get('Duplicate practitioner number not imported', 0)} flagged rows in the full-registration issue file).",
        f"Missing or invalid issued dates ({full_issues.get('Missing or invalid issued date', 0)} flagged full-registration rows).",
        f"Missing graduation year ({full_issues.get('Missing graduation year', 0)} flagged rows in full-registration review).",
        f"Missing institution details ({full_issues.get('Missing institution', 0)} flagged full-registration rows).",
        f"Provisional issue rows with missing or invalid issued date ({prov_issues.get('Missing or invalid issued date; Issued date status: missing_issued_date', 0)} rows).",
        "Conflicting names, mixed spellings, and legacy date formatting from old records.",
    ]:
        add_bullet(doc, item)

    doc.add_paragraph(
        "In simple terms, the platform is increasingly strong, but the reliability of the final register still depends heavily on the quality of source records coming from paper files and old spreadsheets."
    )

    doc.add_heading("9. What must be done for clean data flow from paper to electronic form", level=1)
    for item in [
        "Use one standard intake checklist for every paper submission.",
        "Assign a unique intake reference as soon as the file is received.",
        "Scan the complete paper file early in the process.",
        "Capture records first in a review stage before moving them into the final live register.",
        "Make critical fields mandatory before approval.",
        "Use dropdown lists for standard items such as province, institution, cadre, and document type wherever possible.",
        "Run duplicate checks before approval.",
        "Require registrar verification against the scanned source file.",
        "Use missing-data review tools routinely.",
        "Carry out monthly or quarterly quality audits rather than waiting for one major clean-up cycle.",
    ]:
        add_number(doc, item)

    doc.add_heading("10. Overall conclusion", level=1)
    for text in [
        f"{PROJECT_NAME} is already capable of supporting a broad range of registration, document, dashboard, and reporting functions.",
        "The current live Nursing Council statistics show that the platform is now holding substantial electronic registry data and is able to present it in a readable management format.",
        "The next major gains will come not from adding more screens alone, but from continuously improving data quality, updating employment records, and keeping source records disciplined and standardised.",
    ]:
        doc.add_paragraph(text)

    doc.save(LOCAL_OUTPUT_PATH)
    EXTERNAL_OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    shutil.copyfile(LOCAL_OUTPUT_PATH, EXTERNAL_OUTPUT_PATH)
    return LOCAL_OUTPUT_PATH, EXTERNAL_OUTPUT_PATH


if __name__ == "__main__":
    local_path, external_path = build_document()
    print(local_path)
    print(external_path)
