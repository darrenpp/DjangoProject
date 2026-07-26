from pathlib import Path
import os
import sys

import django
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


BASE_DIR = Path(__file__).resolve().parent.parent
if str(BASE_DIR) not in sys.path:
    sys.path.insert(0, str(BASE_DIR))

os.environ.setdefault("DJANGO_SETTINGS_MODULE", "NDOH_regulatory_bodies.settings")
django.setup()

from apps.dashboard.regulatory_alignment import build_nursing_regulatory_alignment_context  # noqa: E402


DOCS_DIR = BASE_DIR / "docs"
OUTPUT_PATH = DOCS_DIR / "PNG_Nursing_Council_Regulatory_Alignment_Report.docx"
SOURCE_ONE = Path(r"c:\Users\timhi\OneDrive\Desktop\ParotOs\NDOH_Database\Briefs\march_briefs_2026\Final Report-WORKFORCE DATA FOR MINISTER.docx.pdf")
SOURCE_TWO = Path(r"c:\Users\timhi\OneDrive\Desktop\ParotOs\NDOH_Database\Briefs\march_briefs_2026\Situational Analysis report.docx.pdf")


def set_default_font(document: Document):
    style = document.styles["Normal"]
    style.font.name = "Times New Roman"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    style.font.size = Pt(11)

    for style_name in ["Title", "Heading 1", "Heading 2", "Heading 3"]:
        if style_name in document.styles:
            style = document.styles[style_name]
            style.font.name = "Times New Roman"
            style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")


def add_bottom_border(paragraph):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "18")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "000000")
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def add_bullet(document: Document, text: str):
    paragraph = document.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.add_run(text)


def add_table(document: Document, headers, rows, heading=None):
    if heading:
        document.add_paragraph().add_run(heading).bold = True
    table = document.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    header_cells = table.rows[0].cells
    for index, header in enumerate(headers):
        header_cells[index].text = str(header)
        for run in header_cells[index].paragraphs[0].runs:
            run.bold = True

    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cells[index].text = "" if value is None else str(value)
    document.add_paragraph()


def build_document():
    context = build_nursing_regulatory_alignment_context()

    document = Document()
    set_default_font(document)

    section = document.sections[0]
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)

    crest = BASE_DIR / "static" / "img" / "NDOH_LOGO.png"
    if crest.exists():
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.add_run().add_picture(str(crest), width=Inches(1.8))

    heading = document.add_paragraph()
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading_run = heading.add_run("PAPUA NEW GUINEA NURSING COUNCIL\nOFFICE OF THE REGISTRAR")
    heading_run.bold = True
    heading_run.font.size = Pt(15)

    border = document.add_paragraph()
    border.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_bottom_border(border)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("REGULATORY ALIGNMENT AND DATABASE COMPARISON REPORT")
    title_run.bold = True
    title_run.font.size = Pt(14)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run(
        "Comparison of the January 2026 ministerial workforce brief, situational analysis report, "
        "and the current live Nursing Council database."
    )

    document.add_heading("1. Purpose", level=1)
    document.add_paragraph(
        "This report analyses the source documents provided by the PNG Nursing Council, compares their tables and "
        "regulatory statements with the current live database, and records what has now been aligned inside the platform."
    )

    document.add_heading("2. Source Documents Reviewed", level=1)
    add_bullet(document, f"{SOURCE_ONE.name}")
    add_bullet(document, f"{SOURCE_TWO.name}")
    add_bullet(document, f"Live database snapshot generated on {context['generated_on']}")

    document.add_heading("3. Statutory Context and Mandate", level=1)
    document.add_paragraph(context["statutory_context"]["summary"])
    document.add_paragraph().add_run("Core mandate").bold = True
    for item in context["statutory_context"]["mandate_points"]:
        add_bullet(document, item)
    document.add_paragraph().add_run("Outside direct Council mandate").bold = True
    for item in context["statutory_context"]["out_of_scope_points"]:
        add_bullet(document, item)
    document.add_paragraph(context["statutory_context"]["alignment_note"])

    document.add_heading("4. Current Live Registry Snapshot", level=1)
    add_table(document, ["Registry Category", "Current Count", "Active Count"], context["live_snapshot_rows"])
    if context["latest_batch_row"]:
        add_table(
            document,
            ["Latest Source Metric", "Value"],
            [
                ["Source File", context["latest_batch_row"]["source_file_name"]],
                ["Source Kind", context["latest_batch_row"]["source_kind"]],
                ["Completed At", context["latest_batch_row"]["completed_at"]],
                ["Processed Rows", context["latest_batch_row"]["processed_rows"]],
                ["Total Rows", context["latest_batch_row"]["total_rows"]],
            ],
            heading="Latest Imported Source Snapshot",
        )

    document.add_heading("5. High-Level Comparison Summary", level=1)
    add_table(document, context["comparison_summary_headers"], context["comparison_summary_rows"])

    document.add_heading("6. Detailed Table-By-Table Comparison", level=1)
    for table in context["reference_tables"]:
        document.add_heading(table["title"], level=2)
        document.add_paragraph(table["description"])
        document.add_paragraph(f"Source note: {table['source_note']}")
        add_table(document, table["comparison_headers"], table["comparison_rows"], heading="Live Database Comparison")

    document.add_heading("7. Situational Analysis Alignment", level=1)
    document.add_paragraph("The situational analysis was aligned against the platform in three ways: gap analysis, risk alignment, and roadmap readiness.")
    add_table(document, context["gap_analysis_headers"], context["gap_analysis_rows"], heading="Gap Analysis Alignment")
    add_table(document, context["risk_headers"], context["risk_rows"], heading="Risk Alignment")
    add_table(document, context["roadmap_headers"], context["roadmap_rows"], heading="Roadmap Alignment")

    document.add_heading("8. Key Findings", level=1)
    findings = [
        "The live platform is strong for analytics snapshots, imported history, dashboards, privacy scoping, document governance, ICMS case tracking, discipline workflow, decision records, and digital reporting.",
        "The ministerial graduate output tables can be compared meaningfully against qualification records, and several totals are reasonably close rather than exact matches.",
        "The full-licence and registration-elements tables do not align cleanly with current digital records, which indicates either different counting logic, incomplete year capture, or data-cleaning issues between source reports and the live registry.",
        "Employment-type reporting remains a major digital gap because there are currently no captured EmploymentRecord rows in the live system.",
        "The situational analysis remains valid: legislation, approved SOP content, staff training, and disciplined operational use still require management and legal follow-through beyond software alone.",
    ]
    for item in findings:
        add_bullet(document, item)

    document.add_heading("9. Platform Changes Completed", level=1)
    for item in [
        "Added a Nursing Council regulatory intelligence section to the registrar profile.",
        "Added the statutory context and mandate of the PNG Nursing Council into the platform.",
        "Added report-reference tables from the ministerial submission and live database comparison tables beneath them.",
        "Added situational-analysis alignment tables for SWOT, risks, gaps, and roadmap readiness.",
        "Added Nursing Council analytics snapshot reporting from the cleansed workbook.",
        "Added formal ICMS complaints, disciplinary case workflow, and regulatory decision register.",
        "Added document approval/rejection sign-off for controlled repository versions.",
        "Added NHWA workbook alignment as a reporting layer.",
        "Added public FAQ, moderated forum, and mapped institution/facility reference pages.",
        "Added receipt-owner linking and high-value review routing.",
    ]:
        add_bullet(document, item)

    document.add_heading("10. Recommendations", level=1)
    for item in [
        "Standardise legal wording on the Council's statutory basis across all official documents and templates.",
        "Agree a single counting method for graduate outputs, full licences, ATP, and specialisations before using these figures for ministerial or board reporting.",
        "Populate employment and workforce movement data if employment-type reporting is expected from the system.",
        "Adopt SOPs and staff training for formal complaints, discipline, regulatory decisions, document sign-off, and receipt review workflows.",
        "Verify mapped entity coordinates before public map demonstrations.",
        "Use NHWA workbooks as reporting/sign-off outputs only, not as registry overwrite sources.",
        "Use the new registrar profile section as the live operational reference point, but continue validating report-source spreadsheets before executive submission.",
    ]:
        add_bullet(document, item)

    document.save(OUTPUT_PATH)
    return OUTPUT_PATH


if __name__ == "__main__":
    path = build_document()
    print(path)
