import os
import shutil
import subprocess
import sys
import tempfile
import textwrap
import time
import urllib.request
from datetime import datetime
from pathlib import Path

BASE_DIR = Path(__file__).resolve().parents[1]
if str(BASE_DIR) not in sys.path:
    sys.path.insert(0, str(BASE_DIR))

os.environ.setdefault("DJANGO_SETTINGS_MODULE", "NDOH_regulatory_bodies.settings")

import django

django.setup()

from django.contrib.auth import get_user_model
from django.db.models import Count, Sum
from django.test import Client
from django.utils import timezone
from docx import Document as WordDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from PIL import Image as PILImage, ImageDraw, ImageFont
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import cm
from reportlab.platypus import Image as PdfImage
from reportlab.platypus import PageBreak, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

from apps.common.models import DuplicateReviewQueue
from apps.dashboard.ai_provider import ai_provider_status
from apps.complaints.models import ComplaintCase, DisciplinaryCase, RegulatoryDecisionRecord
from apps.dashboard.models import (
    FAQEntry,
    ForumTopic,
    MappedEntity,
    NursingAnalyticsSnapshot,
    NursingLifecycleFact,
    NursingPractitionerIndex,
    Receipt,
)
from apps.dashboard.reference_breakdown import build_reference_breakdown
from apps.dashboard.reports import build_financial_forecast_payload
from apps.documents.models import Document, DocumentApproval, DocumentAuditEvent, DocumentFolder, DocumentVersion
from apps.workforce.models import (
    Application,
    ApplicationPathway,
    CommunityHealthWorker,
    DataImportBatch,
    DocumentRequirement,
    DynamicFormDefinition,
    HealthStudent,
    ImportedWorkbookSheet,
    MedicalDoctor,
    MissingDataReview,
    Midwife,
    NurseAide,
    NursingProfessional,
    PracticingLicenseRecord,
    ProfessionalDocument,
    Qualification,
)


DOCS_DIR = BASE_DIR / "docs"
OUTPUT_DIR = DOCS_DIR / "presentation"
ASSET_DIR = OUTPUT_DIR / "assets"
SCREENSHOT_DIR = ASSET_DIR / "screenshots"
DIAGRAM_DIR = ASSET_DIR / "diagrams"
HTML_DIR = ASSET_DIR / "html"
LOGO_PATH = BASE_DIR / "static" / "img" / "NDOH_LOGO.png"
PROJECT_TITLE = "PNG Nursing Council and Medical Board Online Regulatory Workforce Platform"
DATE_STAMP = "20260601"
DISPLAY_DATE = "1 June 2026"

PACK_MD = OUTPUT_DIR / f"NDOH_Regulatory_Platform_Presentation_Pack_{DATE_STAMP}.md"
PACK_PDF = OUTPUT_DIR / f"NDOH_Regulatory_Platform_Presentation_Pack_{DATE_STAMP}.pdf"
PACK_DOCX = OUTPUT_DIR / f"NDOH_Regulatory_Platform_Presentation_Brief_{DATE_STAMP}.docx"
INDEX_PDF = OUTPUT_DIR / f"NDOH_Regulatory_Platform_Documentation_Index_{DATE_STAMP}.pdf"
README = OUTPUT_DIR / "README.md"


def _ensure_dirs():
    for path in (OUTPUT_DIR, ASSET_DIR, SCREENSHOT_DIR, DIAGRAM_DIR, HTML_DIR):
        path.mkdir(parents=True, exist_ok=True)


def _font(size=22, bold=False):
    candidates = [
        Path("C:/Windows/Fonts/arialbd.ttf" if bold else "C:/Windows/Fonts/arial.ttf"),
        Path("C:/Windows/Fonts/calibrib.ttf" if bold else "C:/Windows/Fonts/calibri.ttf"),
    ]
    for candidate in candidates:
        if candidate.exists():
            return ImageFont.truetype(str(candidate), size)
    return ImageFont.load_default()


def _wrap_text(draw, text, font, max_width):
    words = str(text).split()
    lines = []
    current = ""
    for word in words:
        candidate = f"{current} {word}".strip()
        width = draw.textbbox((0, 0), candidate, font=font)[2]
        if width <= max_width or not current:
            current = candidate
        else:
            lines.append(current)
            current = word
    if current:
        lines.append(current)
    return lines


def _draw_wrapped(draw, xy, text, font, fill, max_width, line_gap=6):
    x, y = xy
    for line in _wrap_text(draw, text, font, max_width):
        draw.text((x, y), line, font=font, fill=fill)
        y += font.size + line_gap
    return y


def _draw_node(draw, box, title, body, fill, outline="#cbd5e1", title_fill="#ffffff"):
    x1, y1, x2, y2 = box
    draw.rounded_rectangle(box, radius=18, fill=fill, outline=outline, width=2)
    draw.text((x1 + 18, y1 + 14), title, font=_font(23, True), fill=title_fill)
    _draw_wrapped(draw, (x1 + 18, y1 + 52), body, _font(17), "#f8fafc" if title_fill == "#ffffff" else "#0f172a", x2 - x1 - 36)


def _arrow(draw, start, end, fill="#334155"):
    draw.line([start, end], fill=fill, width=4)
    ex, ey = end
    sx, sy = start
    if abs(ex - sx) >= abs(ey - sy):
        direction = 1 if ex > sx else -1
        points = [(ex, ey), (ex - 16 * direction, ey - 9), (ex - 16 * direction, ey + 9)]
    else:
        direction = 1 if ey > sy else -1
        points = [(ex, ey), (ex - 9, ey - 16 * direction), (ex + 9, ey - 16 * direction)]
    draw.polygon(points, fill=fill)


def create_architecture_diagram(path):
    image = PILImage.new("RGB", (1800, 1120), "#f8fafc")
    draw = ImageDraw.Draw(image)
    draw.text((70, 50), "Whole Department Enterprise Architecture", font=_font(46, True), fill="#0f172a")
    draw.text((72, 108), "How users, records, reports, AI, backup, and future hosting connect.", font=_font(23), fill="#475569")

    user_nodes = [
        ("Public Users", "Forms, public register search, enquiries", 70, 190),
        ("Professionals", "Own profile, own applications, receipts, documents", 70, 340),
        ("Registrars", "Official review, approvals, licence decisions", 70, 490),
        ("Reviewers", "Assigned reviews after approval", 70, 640),
        ("Finance Users", "Separated financial forecast views", 70, 790),
        ("System Admin", "Security, users, configuration, backups", 70, 940),
    ]
    for title, body, x, y in user_nodes:
        _draw_node(draw, (x, y, x + 380, y + 105), title, body, "#12324a")
        _arrow(draw, (x + 380, y + 52), (590, 560), "#64748b")

    _draw_node(draw, (610, 435, 1030, 685), "Secure Web Platform", "Role-based dashboards, public intake, workforce registry, staff inbox, AI assistant, records hub", "#0f766e")
    service_nodes = [
        ("Workflow Engine", "Application -> checklist -> payment -> review -> decision", 1160, 160, "#1d4f7a"),
        ("Database", "Practitioners, applications, licences, ATP, receipts, source batches", 1160, 310, "#17634f"),
        ("Document Repository", "OpenKM-style official records, versions, OCR/search, metadata", 1160, 460, "#7c2d12"),
        ("Reporting Engine", "Monthly, yearly, financial, Minister and Registrar packs", 1160, 610, "#92400e"),
        ("AI Assistant", "Free local GPT optional; local rules fallback; staff only", 1160, 760, "#334155"),
        ("Backup / Future Hosting", "Government hosting, backup server, restore drills, email alerts", 1160, 910, "#581c87"),
    ]
    for title, body, x, y, fill in service_nodes:
        _draw_node(draw, (x, y, x + 560, y + 112), title, body, fill)
        _arrow(draw, (1030, 560), (x, y + 56), "#64748b")

    image.save(path)


def create_workflow_diagram(path):
    image = PILImage.new("RGB", (1800, 980), "#ffffff")
    draw = ImageDraw.Draw(image)
    draw.text((70, 55), "Regulatory Workflow", font=_font(46, True), fill="#0f172a")
    draw.text((72, 112), "No silent approvals: every official decision needs evidence, status history, and audit trail.", font=_font(23), fill="#475569")
    steps = [
        ("1. Application", "Applicant selects the correct pathway and submits details."),
        ("2. Checklist", "Required documents are generated by pathway."),
        ("3. Payment", "Receipt is uploaded and verified or officially waived."),
        ("4. Review", "Reviewer checks identity, qualifications, competency, and duplicates."),
        ("5. Registrar Decision", "Registrar approves, rejects, or requests missing information."),
        ("6. Licence/Register", "Approved records update the live register and reports."),
    ]
    x, y = 80, 230
    for index, (title, body) in enumerate(steps):
        fill = ["#12324a", "#0f766e", "#92400e", "#1d4f7a", "#7f1d1d", "#17634f"][index]
        _draw_node(draw, (x, y, x + 485, y + 150), title, body, fill)
        if index % 3 != 2:
            _arrow(draw, (x + 485, y + 75), (x + 575, y + 75))
            x += 610
        else:
            x = 80
            y += 270
            if index < len(steps) - 1:
                _arrow(draw, (x + 1510, y - 120), (x + 1510, y - 10))
    image.save(path)


def create_data_flow_diagram(path):
    image = PILImage.new("RGB", (1800, 980), "#f8fafc")
    draw = ImageDraw.Draw(image)
    draw.text((70, 55), "Data Governance Flow", font=_font(46, True), fill="#0f172a")
    draw.text((72, 112), "Imported rows are not automatically trusted. They are staged, validated, cleansed, reviewed, approved, then promoted.", font=_font(23), fill="#475569")
    steps = [
        ("Paper / Excel Source", "Raw files, receipts, ATP lists, application forms"),
        ("Staging Area", "Rows are loaded for inspection, not trusted yet"),
        ("Validation", "Required fields, dates, provinces, identifiers"),
        ("Cleansing", "Aliases, duplicates, missing fields, future dates"),
        ("Human Review", "Registrar, data-quality, finance or reviewer decision"),
        ("Live Registry", "Approved people, licences, qualifications, payments"),
        ("Reports", "Dashboards, monthly packs, Minister brief, public-safe outputs"),
    ]
    x = 70
    y = 360
    for index, (title, body) in enumerate(steps):
        width = 225
        fill = "#12324a" if index in {0, 5, 6} else "#0f766e" if index in {1, 2, 3} else "#92400e"
        _draw_node(draw, (x, y, x + width, y + 180), title, body, fill)
        if index < len(steps) - 1:
            _arrow(draw, (x + width, y + 90), (x + width + 45, y + 90))
        x += width + 70
    draw.rounded_rectangle((72, 690, 1725, 835), radius=18, fill="#e0f2fe", outline="#7dd3fc", width=2)
    _draw_wrapped(
        draw,
        (100, 720),
        "Plain explanation: old spreadsheets and paper files can contain spelling differences, missing dates, duplicate people, wrong provinces, and future dates. The platform makes those issues visible before figures are used for management decisions.",
        _font(25, True),
        "#0f172a",
        1580,
    )
    image.save(path)


def create_role_diagram(path):
    image = PILImage.new("RGB", (1800, 980), "#ffffff")
    draw = ImageDraw.Draw(image)
    draw.text((70, 55), "Role Access and Privacy Model", font=_font(46, True), fill="#0f172a")
    draw.text((72, 112), "Each user group sees only the work and records needed for its role.", font=_font(23), fill="#475569")
    nodes = [
        ("Public / Applicant", "Own forms and public-safe register only", 80, 230, "#475569"),
        ("Nursing Council", "Nursing, midwifery, nurse aide, graduand records", 500, 230, "#0f766e"),
        ("Medical Board", "Doctor, CHW, Medical Board records", 920, 230, "#1d4f7a"),
        ("Finance", "Financial forecast and receipt summaries only", 1340, 230, "#92400e"),
        ("Data Quality", "Missing fields, duplicates, source corrections", 290, 550, "#7c2d12"),
        ("System Admin", "Configuration, security, backups, users", 710, 550, "#12324a"),
        ("Document Repository", "Office-scoped records, versions, audit history", 1130, 550, "#581c87"),
    ]
    for title, body, x, y, fill in nodes:
        _draw_node(draw, (x, y, x + 370, y + 145), title, body, fill)
    draw.rounded_rectangle((80, 805, 1720, 900), radius=18, fill="#fef3c7", outline="#f59e0b", width=2)
    _draw_wrapped(draw, (112, 830), "Privacy rule: Nursing Council and Medical Board data must remain separated through backend permissions, not only hidden buttons.", _font(28, True), "#111827", 1580)
    image.save(path)


def create_diagrams():
    diagrams = {
        "architecture": DIAGRAM_DIR / "enterprise_architecture.png",
        "workflow": DIAGRAM_DIR / "regulatory_workflow.png",
        "data_flow": DIAGRAM_DIR / "data_governance_flow.png",
        "roles": DIAGRAM_DIR / "role_access_privacy.png",
    }
    create_architecture_diagram(diagrams["architecture"])
    create_workflow_diagram(diagrams["workflow"])
    create_data_flow_diagram(diagrams["data_flow"])
    create_role_diagram(diagrams["roles"])
    return diagrams


def _first_user(role, department_contains=None):
    User = get_user_model()
    queryset = User.objects.filter(role=role, is_active=True)
    if department_contains:
        queryset = queryset.filter(department__icontains=department_contains)
    user = queryset.order_by("id").first()
    if user:
        return user

    username = f"presentation_{role}_{department_contains or 'general'}".replace(" ", "_").lower()
    user = User(username=username, role=role, department=department_contains or "")
    user.set_unusable_password()
    user.email = f"{username}@example.local"
    user.is_active = True
    user.role_approved = True
    user.operations_approved = True
    if role in {"admin", "registrar", "reviewer"}:
        user.is_staff = True
    if role == "admin":
        user.is_superuser = True
    user.save()
    return user


def _presentation_users():
    return {
        "admin": _first_user("admin"),
        "nursing_registrar": _first_user("registrar", "Nursing"),
        "medical_registrar": _first_user("registrar", "Medical"),
        "finance": _first_user("reviewer", "Finance"),
        "data_quality": _first_user("reviewer", "Data Quality"),
        "nurse": _first_user("nurse"),
        "doctor": _first_user("doctor"),
        "graduand": _first_user("graduand"),
    }


def _start_static_server(port):
    manage_py = BASE_DIR / "manage.py"
    process = subprocess.Popen(
        [sys.executable, str(manage_py), "runserver", f"127.0.0.1:{port}", "--noreload"],
        cwd=str(BASE_DIR),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
    )
    for _ in range(30):
        try:
            urllib.request.urlopen(f"http://127.0.0.1:{port}/", timeout=1)
            return process
        except Exception:
            time.sleep(0.4)
    process.terminate()
    return None


def _chrome_binary():
    for candidate in [
        shutil.which("chrome.exe"),
        shutil.which("chrome"),
        r"C:\Program Files\Google\Chrome\Application\chrome.exe",
        r"C:\Program Files (x86)\Microsoft\Edge\Application\msedge.exe",
        shutil.which("msedge.exe"),
    ]:
        if candidate and Path(candidate).exists():
            return str(candidate)
    return ""


def _inject_capture_head(html, port):
    head = (
        f'<base href="http://127.0.0.1:{port}/">'
        "<style>"
        "html,body{background:#f8fafc!important;color:#0f172a!important;}"
        "*{text-shadow:none!important;}"
        ".content-wrapper,.card,.card-body,.table td,.table th{color:#0f172a!important;}"
        ".bg-primary *,.bg-success *,.bg-info *,.bg-danger *,.bg-dark *,.bg-secondary *{color:#fff!important;}"
        ".bg-warning *,.bg-light *,.bg-white *{color:#111827!important;}"
        ".main-sidebar,.main-header,.main-footer{display:none!important;}"
        ".app-wrapper,.wrapper{display:block!important;width:100%!important;min-height:auto!important;}"
        ".content-wrapper{margin:0!important;width:100%!important;min-height:1100px!important;position:static!important;transform:none!important;}"
        ".content-wrapper>.content,.content,.content>.container-fluid,.container-fluid{max-width:none!important;width:100%!important;margin:0!important;padding:18px!important;}"
        ".content-header{padding:18px 18px 0!important;}"
        ".row{max-width:100%!important;}"
        ".helpdesk-launcher,.helpdesk-widget{display:none!important;}"
        "</style>"
    )
    if "<head>" in html:
        return html.replace("<head>", f"<head>{head}", 1)
    return head + html


def _clean_generated_html(html):
    """Keep captured HTML snapshots stable enough for git diff checks."""
    return "\n".join(line.rstrip() for line in html.splitlines()) + "\n"


def _fallback_screenshot(path, title, description):
    image = PILImage.new("RGB", (1600, 950), "#f8fafc")
    draw = ImageDraw.Draw(image)
    draw.rounded_rectangle((70, 70, 1530, 880), radius=24, fill="#ffffff", outline="#cbd5e1", width=2)
    if LOGO_PATH.exists():
        logo = PILImage.open(LOGO_PATH).convert("RGBA")
        logo.thumbnail((150, 110))
        image.paste(logo, (110, 105), logo)
    draw.text((290, 120), PROJECT_TITLE, font=_font(36, True), fill="#0f172a")
    draw.text((290, 185), title, font=_font(44, True), fill="#0f766e")
    _draw_wrapped(draw, (115, 290), description, _font(28), "#334155", 1370)
    draw.text((115, 770), "Screenshot capture fallback generated for presentation pack.", font=_font(22), fill="#64748b")
    image.save(path)


def _json_table_screenshot(path, title, description, payload):
    image = PILImage.new("RGB", (1700, 1050), "#f8fafc")
    draw = ImageDraw.Draw(image)
    draw.rounded_rectangle((55, 50, 1645, 990), radius=24, fill="#ffffff", outline="#cbd5e1", width=2)
    if LOGO_PATH.exists():
        logo = PILImage.open(LOGO_PATH).convert("RGBA")
        logo.thumbnail((120, 90))
        image.paste(logo, (95, 78), logo)
    draw.text((240, 80), title, font=_font(42, True), fill="#0f172a")
    _draw_wrapped(draw, (240, 135), description, _font(21), "#475569", 1300)
    count = payload.get("count", len(payload.get("results", []))) if isinstance(payload, dict) else 0
    draw.text((95, 225), f"Public-safe results returned: {count}", font=_font(28, True), fill="#0f766e")
    headers = ["Full name", "Registration", "Category", "Licence status", "Eligible"]
    col_x = [95, 540, 805, 1080, 1355]
    widths = [420, 240, 250, 250, 190]
    y = 300
    draw.rounded_rectangle((80, y - 16, 1585, y + 48), radius=10, fill="#0f766e")
    for x, header in zip(col_x, headers):
        draw.text((x, y), header, font=_font(19, True), fill="#ffffff")
    y += 75
    results = payload.get("results", [])[:9] if isinstance(payload, dict) else []
    for index, row in enumerate(results):
        fill = "#f8fafc" if index % 2 else "#eef6f5"
        draw.rounded_rectangle((80, y - 14, 1585, y + 50), radius=8, fill=fill)
        values = [
            row.get("full_name", ""),
            row.get("registration_number", ""),
            row.get("professional_category", ""),
            row.get("licence_status", ""),
            "Yes" if row.get("eligible_to_practice") else "No",
        ]
        for x, width, value in zip(col_x, widths, values):
            _draw_wrapped(draw, (x, y), str(value), _font(18), "#0f172a", width, line_gap=2)
        y += 72
    draw.text((95, 940), "Privacy note: public search shows safe fields only. It does not expose DOB, phone, address, passport, medical, police, transcript, receipt image, or internal notes.", font=_font(19), fill="#475569")
    image.save(path)


def capture_screenshots():
    users = _presentation_users()
    pages = [
        ("public_home", "Public Home Page", "/", None, "Public entry point for applicants and general users."),
        ("login", "Login Page", "/accounts/login/", None, "Secure sign-in page for staff and registered users."),
        ("overall_dashboard", "Overall Dashboard", "/dashboard/", users["admin"], "Leadership dashboard with current workforce summary."),
        ("production_readiness", "Production Readiness Dashboard", "/dashboard/production-readiness/", users["admin"], "Launch-readiness view showing remaining data and governance work."),
        ("nursing_council_dashboard", "Nursing Council Registrar Dashboard", "/dashboard/nursing-council/", users["nursing_registrar"], "Nursing Council workspace for registry, ATP, forms, and operations."),
        ("medical_board_dashboard", "Medical Board Dashboard", "/dashboard/medical-board/", users["medical_registrar"], "Medical Board workspace, kept separate from Nursing Council records."),
        ("workforce_flow", "Workforce Flow", "/dashboard/flow/", users["nursing_registrar"], "Workforce movement and planning dashboard."),
        ("financial_forecast", "Financial Forecast", "/dashboard/reports/financial/?office=nursing", users["finance"], "Separated financial forecast page for receipt and revenue tracking."),
        ("staff_ai", "Staff AI Assistant", "/dashboard/staff-ai/", users["nursing_registrar"], "Staff-only assistant for operational questions and data-quality guidance."),
        ("documents_search", "Document Repository Search", "/documents/search/", users["nursing_registrar"], "OpenKM-style repository search and records management entry point."),
        ("complaints_icms", "ICMS Complaints Register", "/dashboard/complaints/", users["nursing_registrar"], "Formal complaint, incident, and enquiry case-management register."),
        ("disciplinary_cases", "Disciplinary Case Register", "/dashboard/complaints/discipline/", users["nursing_registrar"], "Disciplinary pathway tracking with stages, events, attachments, and escalation."),
        ("decision_register", "Regulatory Decision Register", "/dashboard/complaints/decisions/", users["nursing_registrar"], "Formal decision records with rationale, authority, evidence, conditions, and appeal rights."),
        ("nhwa_workbooks", "NHWA Workbook Centre", "/dashboard/nhwa-workbooks/", users["admin"], "NHWA standards and reporting workbook layer populated from verified platform data."),
        ("public_faqs", "Public FAQs", "/dashboard/public/faqs/", None, "Public frequently asked questions and guidance."),
        ("public_forum", "Public And Practitioner Forum", "/dashboard/public/forum/", None, "Moderated public and role-based discussion categories."),
        ("public_map", "Mapped Schools, Institutions, and Facilities", "/dashboard/public/map/", None, "Mapped reference entities using locally stored verified coordinates."),
        ("records_hub", "Records Hub", "/records/", users["nursing_registrar"], "Staff record management hub for authorised operational users."),
        ("nursing_professionals", "Nursing Professionals Records Table", "/records/nursingprofessional/", users["nursing_registrar"], "Registrar records table with search, sorting, pagination, and authorised CRUD actions."),
        ("duplicate_review_queue", "Duplicate Review Queue", "/dashboard/duplicate-reviews/", users["nursing_registrar"], "Data-quality duplicate queue with grouped source rows and review actions."),
        ("staff_notifications", "Staff Inbox and Notifications", "/notifications/communications/", users["nursing_registrar"], "Staff mailbox, notification history, read/opened status, and access requests."),
        ("nursing_forms", "Nursing Forms Portal", "/nursing/forms/", None, "Public Nursing Council pathway and form selection screen."),
        ("public_register", "Public Nursing Register Search", "/public/nursing-council/register/search/", None, "Safe public register search with limited fields."),
        ("nurse_portal", "Nurse User Portal", "/dashboard/nurse/", users["nurse"], "Individual nurse self-service dashboard."),
        ("graduand_portal", "Graduand User Portal", "/dashboard/student/", users["graduand"], "Graduand pathway dashboard."),
        ("doctor_portal", "Doctor User Portal", "/dashboard/doctor/", users["doctor"], "Medical professional self-service dashboard."),
    ]
    chrome = _chrome_binary()
    port = 8765
    server = _start_static_server(port)
    screenshots = []
    try:
        for slug, title, path, user, description in pages:
            output_png = SCREENSHOT_DIR / f"{slug}.png"
            html_file = HTML_DIR / f"{slug}.html"
            try:
                client = Client(HTTP_HOST=f"127.0.0.1:{port}")
                if user is not None:
                    client.force_login(user)
                response = client.get(path, follow=True)
                html = response.content.decode("utf-8", errors="replace")
                content_type = response.headers.get("Content-Type", "")
                if "application/json" in content_type:
                    try:
                        _json_table_screenshot(output_png, title, description, response.json())
                    except Exception:
                        _fallback_screenshot(output_png, title, description)
                elif chrome and server:
                    html = _inject_capture_head(html, port)
                    html_file.write_text(_clean_generated_html(html), encoding="utf-8")
                    with tempfile.TemporaryDirectory() as user_data_dir:
                        result = subprocess.run(
                            [
                                chrome,
                                "--headless=new",
                                "--disable-gpu",
                                "--disable-extensions",
                                "--hide-scrollbars",
                                "--allow-file-access-from-files",
                                f"--user-data-dir={user_data_dir}",
                                "--window-size=1700,1050",
                                "--virtual-time-budget=2500",
                                f"--screenshot={output_png}",
                                html_file.as_uri(),
                            ],
                            cwd=str(BASE_DIR),
                            stdout=subprocess.DEVNULL,
                            stderr=subprocess.DEVNULL,
                            timeout=25,
                        )
                    if result.returncode != 0 or not output_png.exists():
                        _fallback_screenshot(output_png, title, description)
                else:
                    html = _inject_capture_head(html, port)
                    html_file.write_text(_clean_generated_html(html), encoding="utf-8")
                    _fallback_screenshot(output_png, title, description)
            except Exception as exc:
                _fallback_screenshot(output_png, title, f"{description} Capture note: {exc}")
            screenshots.append({"slug": slug, "title": title, "path": output_png, "description": description, "url": path})
    finally:
        if server:
            server.terminate()
    return screenshots


def _safe_count(model):
    try:
        return model.objects.count()
    except Exception:
        return 0


def collect_live_statistics():
    reference = build_reference_breakdown()
    latest_batch = DataImportBatch.objects.order_by("-completed_at", "-started_at").first()
    active_snapshot = NursingAnalyticsSnapshot.objects.filter(is_active=True).order_by("-activated_at", "-created_at").first()
    analytics_kpis = active_snapshot.kpi_summary if active_snapshot else {}
    application_status_rows = list(
        Application.objects.values("status").annotate(total=Count("id")).order_by("status")
    )
    imported_record_mix = list(
        PracticingLicenseRecord.objects.values("record_type").annotate(total=Count("id")).order_by("record_type")
    )
    receipt_summary = Receipt.objects.aggregate(total=Count("id"), amount=Sum("amount"))
    latest_sheet = ImportedWorkbookSheet.objects.order_by("-id").first()
    finance = build_financial_forecast_payload("nursing", generated_by="Presentation package")
    nursing_finance = finance["offices"].get("nursing", {})
    nursing_cleaned_totals = [
        (
            "Nursing Council total lifecycle records",
            analytics_kpis.get("total_lifecycle_records", NursingLifecycleFact.objects.filter(snapshot=active_snapshot).count() if active_snapshot else 0),
            "Cleansed analytics snapshot rows across provisional, full licence, and ATP lifecycle stages.",
        ),
        (
            "Clean ATP records",
            analytics_kpis.get("clean_atp_records", 0),
            "Cleansed Authority to Practice records from the active Nursing Council analytics snapshot.",
        ),
        (
            "Clean provisional records",
            analytics_kpis.get("clean_provisional_records", 0),
            "Cleansed provisional licence records from the active Nursing Council analytics snapshot.",
        ),
        (
            "Clean full-licence records",
            analytics_kpis.get("clean_full_licence_records", 0),
            "Cleansed full-licence records from the active Nursing Council analytics snapshot.",
        ),
        (
            "Estimated practitioner match groups",
            analytics_kpis.get("estimated_practitioner_match_groups", NursingPractitionerIndex.objects.filter(snapshot=active_snapshot).count() if active_snapshot else 0),
            "Analytics grouping count used for workforce analysis; not a legal practitioner ID.",
        ),
        (
            "Data quality health score",
            f"{analytics_kpis.get('data_quality_health_score', 0)}%",
            "Current cleansed-data quality score for the active Nursing Council analytics snapshot.",
        ),
    ]
    nursing_cadre_rows = []
    if active_snapshot:
        nursing_cadre_rows = list(
            active_snapshot.cadre_stage_metrics
            .order_by("-grand_total", "cadre")
            .values(
                "cadre",
                "provisional_licence_count",
                "full_licence_count",
                "authority_to_practice_count",
                "grand_total",
            )[:14]
        )
    return {
        "generated_on": DISPLAY_DATE,
        "headline": nursing_cleaned_totals,
        "nursing_cadre_rows": nursing_cadre_rows,
        "operational_status": [
            ("Live legal RN person table", _safe_count(NursingProfessional), "Operational NursingProfessional rows only; not the cleansed Nursing Council analytics total."),
            ("Live legal midwife person table", _safe_count(Midwife), "Operational Midwife rows only; not the cleansed Nursing Council analytics total."),
            ("Live legal nurse aide person table", _safe_count(NurseAide), "Operational NurseAide rows only; not the cleansed Nursing Council analytics total."),
            ("Live legal graduand/student table", _safe_count(HealthStudent), "Operational HealthStudent rows only; not the cleansed provisional analytics total."),
            ("Community Health Workers", _safe_count(CommunityHealthWorker), "Medical Board / CHW scope records currently loaded."),
            ("Medical Doctors", _safe_count(MedicalDoctor), "Medical Board doctor records currently loaded."),
            ("Applications", _safe_count(Application), "All application records currently stored."),
            ("Receipts", receipt_summary.get("total") or 0, f"Receipt records currently stored. Total amount: PGK {receipt_summary.get('amount') or 0}."),
            ("Imported Licence / History Rows", _safe_count(PracticingLicenseRecord), "Operational and historical spreadsheet rows; one person can have multiple rows."),
            ("Qualifications", _safe_count(Qualification), "Qualification records currently stored."),
            ("Missing Data Review Items", _safe_count(MissingDataReview), "Data-quality items created for review and correction."),
            ("Pending Missing Data Review Items", MissingDataReview.objects.filter(status="pending").count(), "Missing-data items still pending."),
            ("Duplicate Review Items", _safe_count(DuplicateReviewQueue), "Possible duplicate records requiring staff review."),
            ("Pending Duplicate Review Items", DuplicateReviewQueue.objects.filter(status="pending").count(), "Duplicate-review items still pending."),
        ],
        "reference": reference,
        "active_snapshot": active_snapshot,
        "nursing_analytics": {
            "lifecycle_facts": NursingLifecycleFact.objects.filter(snapshot=active_snapshot).count() if active_snapshot else 0,
            "practitioner_groups": NursingPractitionerIndex.objects.filter(snapshot=active_snapshot).count() if active_snapshot else 0,
            "kpis": analytics_kpis,
            "source_file": active_snapshot.source_file_name if active_snapshot else "-",
            "generated_on": active_snapshot.workbook_generated_on if active_snapshot else "-",
        },
        "latest_batch": latest_batch,
        "latest_sheet": latest_sheet,
        "application_status_rows": application_status_rows,
        "imported_record_mix": imported_record_mix,
        "documents": {
            "folders": _safe_count(DocumentFolder),
            "documents": _safe_count(Document),
            "versions": _safe_count(DocumentVersion),
            "audit_events": _safe_count(DocumentAuditEvent),
            "approvals": _safe_count(DocumentApproval),
            "legacy_uploads": _safe_count(ProfessionalDocument),
        },
        "workflow_config": {
            "pathways": _safe_count(ApplicationPathway),
            "forms": _safe_count(DynamicFormDefinition),
            "requirements": _safe_count(DocumentRequirement),
        },
        "case_management": {
            "complaints": _safe_count(ComplaintCase),
            "disciplinary_cases": _safe_count(DisciplinaryCase),
            "decisions": _safe_count(RegulatoryDecisionRecord),
        },
        "engagement": {
            "mapped_entities": _safe_count(MappedEntity),
            "geocoded_entities": MappedEntity.objects.exclude(latitude__isnull=True).exclude(longitude__isnull=True).count(),
            "faqs": _safe_count(FAQEntry),
            "forum_topics": _safe_count(ForumTopic),
        },
        "finance": nursing_finance,
        "ai": ai_provider_status(),
    }


def _format_number(value):
    try:
        return f"{int(value):,}"
    except Exception:
        return str(value)


def _markdown_table(headers, rows):
    lines = ["| " + " | ".join(headers) + " |", "| " + " | ".join(["---"] * len(headers)) + " |"]
    for row in rows:
        lines.append("| " + " | ".join(str(cell).replace("\n", " ") for cell in row) + " |")
    return "\n".join(lines)


def build_markdown(stats, diagrams, screenshots):
    headline_rows = [(label, _format_number(value), meaning) for label, value, meaning in stats["headline"]]
    operational_rows = [(label, _format_number(value), meaning) for label, value, meaning in stats["operational_status"]]
    cadre_rows = [
        (
            row["cadre"],
            _format_number(row["provisional_licence_count"]),
            _format_number(row["full_licence_count"]),
            _format_number(row["authority_to_practice_count"]),
            _format_number(row["grand_total"]),
        )
        for row in stats["nursing_cadre_rows"]
    ]
    app_rows = [(row["status"].title(), _format_number(row["total"])) for row in stats["application_status_rows"]]
    record_rows = [(row["record_type"].replace("_", " ").title(), _format_number(row["total"])) for row in stats["imported_record_mix"][:12]]
    screenshot_rows = [(item["title"], item["url"], str(item["path"].relative_to(OUTPUT_DIR))) for item in screenshots]
    latest_batch = stats["latest_batch"]
    latest_sheet = stats["latest_sheet"]
    finance = stats["finance"]
    reference = stats["reference"]
    ai = stats["ai"]
    analytics = stats["nursing_analytics"]
    analytics_kpis = analytics.get("kpis", {})
    cases = stats["case_management"]
    engagement = stats["engagement"]
    latest_sheet_rows = "-"
    if latest_sheet:
        latest_sheet_rows = getattr(latest_sheet, "imported_rows", None) or getattr(latest_sheet, "raw_rows", None) or "-"
    md = f"""# Presentation Pack

Project: {PROJECT_TITLE}

Generated: {stats["generated_on"]}

## 1. Plain-Language Overview

This system is a government regulatory operations platform. It helps the National Department of Health, the PNG Nursing Council, and the Medical Board manage applications, practitioner records, licences, qualifications, documents, receipts, imports, data-quality issues, dashboards, reports, and staff workflows.

The simple rule is:

**Imported rows are not automatically trusted. They are staged, validated, cleansed, reviewed, approved, then promoted into live registry records.**

## 2. Nursing Council Cleansed Analytics Totals

{_markdown_table(["Statistic", "Current total", "Meaning"], headline_rows)}

These are the current Nursing Council figures after the cleanse. They come from the active cleansed analytics snapshot, not from the legal live person tables. The legal registry remains protected until records are promoted through approved workflow.

## 2A. Cleansed Nursing Council Cadre / Stage Breakdown

{_markdown_table(["Cadre", "Provisional", "Full licence", "ATP", "Total"], cadre_rows)}

## 3. Source And Recency

- Latest import batch: {latest_batch.source_file_name if latest_batch else "No batch found"}
- Latest import status: {latest_batch.status if latest_batch else "-"}
- Latest import completed at: {latest_batch.completed_at if latest_batch else "-"}
- Latest workbook sheet: {latest_sheet.sheet_name if latest_sheet else "-"}
- Latest workbook rows processed: {latest_sheet_rows}
- Active Nursing Council analytics source: {analytics.get("source_file")}
- Active Nursing Council analytics generated on: {analytics.get("generated_on")}

## 3A. Operational Live Registry And Platform Counts

{_markdown_table(["Statistic", "Current total", "Meaning"], operational_rows)}

`Person_Group_Key` is used for analytics grouping only. It is not a legal practitioner identity.

## 4. Nursing Council Institution And Facility Breakdown

{_markdown_table(["Breakdown item", "Count"], [
    ("Recognised PNG nursing schools", reference["png_nursing_school_count"]),
    ("Government nursing schools", reference["government_nursing_school_count"]),
    ("Non-government nursing schools", reference["non_government_nursing_school_count"]),
    ("Ownership review still needed", reference["review_nursing_school_count"]),
    ("Raw institution rows", reference["raw_institution_total"]),
    ("CHW training references", reference["chw_training_reference_count"]),
    ("Overseas institution references", reference["overseas_institution_reference_count"]),
    ("Local nursing-like names needing cleansing", reference["unmapped_local_nursing_reference_count"]),
    ("Cleaned workplace references from imports", reference["facility_grouped_reference_count"]),
    ("Raw distinct workplace addresses", reference["facility_raw_reference_count"]),
])}

## 5. Application Status Totals

{_markdown_table(["Status", "Current total"], app_rows)}

## 6. Imported Record Activity Mix

{_markdown_table(["Record activity", "Current total"], record_rows)}

## 7. Document Management / OpenKM-Style Repository

{_markdown_table(["Repository item", "Current total"], [
    ("Folders", stats["documents"]["folders"]),
    ("Documents", stats["documents"]["documents"]),
    ("Versions", stats["documents"]["versions"]),
    ("Document audit events", stats["documents"]["audit_events"]),
    ("Document approvals/rejections", stats["documents"]["approvals"]),
    ("Legacy professional uploads", stats["documents"]["legacy_uploads"]),
])}

## 8. Workflow Configuration

{_markdown_table(["Configuration item", "Current total"], [
    ("Application pathways", stats["workflow_config"]["pathways"]),
    ("Dynamic form definitions", stats["workflow_config"]["forms"]),
    ("Document requirements", stats["workflow_config"]["requirements"]),
    ("ICMS complaint cases", cases["complaints"]),
    ("Disciplinary cases", cases["disciplinary_cases"]),
    ("Regulatory decision records", cases["decisions"]),
])}

## 9. Nursing Council Finance Summary

{_markdown_table(["Financial item", "Current value"], [
    ("Manual completed receipts", finance.get("manual_completed_count", 0)),
    ("Manual completed total", f"PGK {finance.get('manual_completed_total', 0)}"),
    ("Spreadsheet receipt rows", finance.get("imported_count", 0)),
    ("Spreadsheet receipt total", f"PGK {finance.get('imported_total', 0)}"),
    ("Combined total", f"PGK {finance.get('combined_total', 0)}"),
    ("Date-quality issues", finance.get("date_quality_issue_count", 0)),
])}

## 9A. Public Engagement And Mapping

{_markdown_table(["Item", "Current total"], [
    ("Mapped entity references", engagement["mapped_entities"]),
    ("Mapped entities with stored coordinates", engagement["geocoded_entities"]),
    ("FAQ entries", engagement["faqs"]),
    ("Forum topics", engagement["forum_topics"]),
])}

Google Maps reads locally stored coordinates. The page does not geocode every load.

## 10. AI Assistant Position

- Current AI mode: {ai["label"]}
- Current AI detail: {ai["detail"]}
- Free local GPT ready: {ai.get("ollama_ready")}
- Model configured: {ai.get("ollama_model") or ai.get("local_model") or "Local rule-based mode"}

The AI assistant is for staff guidance only. It does not approve applications, issue licences, or write imported rows directly into live registry records.

## 11. Diagrams

- Enterprise Architecture: {diagrams["architecture"].relative_to(OUTPUT_DIR)}
- Regulatory Workflow: {diagrams["workflow"].relative_to(OUTPUT_DIR)}
- Data Governance Flow: {diagrams["data_flow"].relative_to(OUTPUT_DIR)}
- Role Access And Privacy: {diagrams["roles"].relative_to(OUTPUT_DIR)}

## 12. Current Interface Screenshots

{_markdown_table(["Screen", "System path", "Screenshot file"], screenshot_rows)}

## 13. Presentation Talking Points

- The platform is already operating as a structured registry, workflow, reporting, finance, and records-management system.
- Nursing Council and Medical Board workspaces are separated to support privacy and proper regulatory governance.
- Board-specific dashboards now show the correct Nursing Council or Medical Board welcome header with PNG emblem identity.
- Public registration now uses controlled Role/Cadre dropdowns, including separate CHW provisional and CHW full-license pathways.
- Notification history, unread badge clearing, and opened/read mailbox status are now part of the operational workflow.
- Records Hub and Duplicate Review Queue now include table search, sorting, page length, and pagination functions for registrar/data-quality work.
- Formal ICMS complaints, discipline workflow, and regulatory decision register now support defensible case management.
- Repository documents can be approved or rejected as controlled current versions.
- Nursing Council analytics are powered by an active cleansed snapshot and server-side drilldowns.
- NHWA workbooks are a reporting layer populated from verified platform data and do not overwrite legal registry records automatically.
- Public FAQs, moderated forums, and mapped institutions/facilities provide a cleaner public and practitioner experience.
- The system now shows clean institution counts separately from raw historical reference rows.
- The biggest ongoing risk is data quality from old paper files and legacy spreadsheets, not the user interface itself.
- Staff should use the Production Readiness dashboard, duplicate review, missing-data review, and import preview tools before publishing management statistics.
- Free local GPT support is available through Ollama or an approved internal model server, with safe fallback to local rules.
"""
    PACK_MD.write_text(md, encoding="utf-8")
    return md


styles = getSampleStyleSheet()
styles.add(ParagraphStyle("CoverTitle", parent=styles["Title"], fontSize=18, leading=22, alignment=TA_CENTER, textColor=colors.HexColor("#12324A")))
styles.add(ParagraphStyle("SectionTitle", parent=styles["Heading1"], fontSize=14, leading=17, textColor=colors.HexColor("#0F766E"), spaceBefore=12, spaceAfter=8))
styles.add(ParagraphStyle("BodyClean", parent=styles["BodyText"], fontSize=9, leading=12))
styles.add(ParagraphStyle("SmallClean", parent=styles["BodyText"], fontSize=7.6, leading=9.5))


def _pdf_footer(canvas, doc):
    canvas.saveState()
    canvas.setFont("Helvetica", 7)
    canvas.setFillColor(colors.HexColor("#12324A"))
    canvas.drawString(1.4 * cm, 1 * cm, "Confidential official presentation pack - NDOH regulatory platform")
    canvas.drawRightString(19.6 * cm, 1 * cm, f"Page {doc.page}")
    canvas.restoreState()


def _pdf_table(headers, rows, widths=None):
    data = [[Paragraph(str(cell), styles["SmallClean"]) for cell in headers]]
    for row in rows:
        data.append([Paragraph(str(cell), styles["SmallClean"]) for cell in row])
    table = Table(data, colWidths=widths, repeatRows=1, hAlign="LEFT")
    table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#0F766E")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("GRID", (0, 0), (-1, -1), 0.25, colors.HexColor("#CBD5E1")),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#F8FAFC")]),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("LEFTPADDING", (0, 0), (-1, -1), 4),
        ("RIGHTPADDING", (0, 0), (-1, -1), 4),
        ("TOPPADDING", (0, 0), (-1, -1), 4),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
    ]))
    return table


def _pdf_image(path, max_width=17.6 * cm, max_height=12.2 * cm):
    with PILImage.open(path) as image:
        width, height = image.size
    ratio = min(max_width / width, max_height / height)
    return PdfImage(str(path), width=width * ratio, height=height * ratio)


def build_presentation_pdf(stats, diagrams, screenshots):
    doc = SimpleDocTemplate(str(PACK_PDF), pagesize=A4, rightMargin=1.25 * cm, leftMargin=1.25 * cm, topMargin=1.25 * cm, bottomMargin=1.55 * cm)
    story = []
    if LOGO_PATH.exists():
        story.append(PdfImage(str(LOGO_PATH), width=2.6 * cm, height=2.2 * cm))
    story.append(Paragraph("PAPUA NEW GUINEA NATIONAL DEPARTMENT OF HEALTH", styles["CoverTitle"]))
    story.append(Paragraph(PROJECT_TITLE, styles["CoverTitle"]))
    story.append(Paragraph("Presentation Pack For Management, Staff Training, ICT Review, And Launch Readiness", styles["BodyClean"]))
    story.append(Paragraph(f"Generated: {stats['generated_on']}", styles["BodyClean"]))
    story.append(PageBreak())

    story.append(Paragraph("1. Executive Overview In Plain Language", styles["SectionTitle"]))
    story.append(Paragraph("The platform is a government regulatory operations system. It manages applications, practitioner records, licences, qualifications, documents, receipts, imports, data-quality issues, dashboards, reports, and staff workflows for the Nursing Council and Medical Board.", styles["BodyClean"]))
    story.append(Paragraph("Core rule: imported rows are not automatically trusted. They are staged, validated, cleansed, reviewed, approved, then promoted into live registry records.", styles["BodyClean"]))
    story.append(Spacer(1, 0.2 * cm))
    story.append(Paragraph("Nursing Council cleansed analytics totals", styles["SectionTitle"]))
    story.append(_pdf_table(["Statistic", "Current total", "Meaning"], [(label, _format_number(value), meaning) for label, value, meaning in stats["headline"]], [4.1 * cm, 3.0 * cm, 10.3 * cm]))
    story.append(Spacer(1, 0.25 * cm))
    story.append(Paragraph("These are the current Nursing Council figures after the cleanse. They come from the active analytics snapshot; legal person-table records remain protected until controlled workflow promotion.", styles["BodyClean"]))
    story.append(Spacer(1, 0.2 * cm))
    story.append(_pdf_table(
        ["Cadre", "Provisional", "Full licence", "ATP", "Total"],
        [
            (
                row["cadre"],
                _format_number(row["provisional_licence_count"]),
                _format_number(row["full_licence_count"]),
                _format_number(row["authority_to_practice_count"]),
                _format_number(row["grand_total"]),
            )
            for row in stats["nursing_cadre_rows"][:8]
        ],
        [5.2 * cm, 2.7 * cm, 2.7 * cm, 2.7 * cm, 2.7 * cm],
    ))

    story.append(PageBreak())
    story.append(Paragraph("2. Department-Wide Diagrams", styles["SectionTitle"]))
    for title, path in [
        ("Enterprise Architecture", diagrams["architecture"]),
        ("Regulatory Workflow", diagrams["workflow"]),
        ("Data Governance Flow", diagrams["data_flow"]),
        ("Role Access And Privacy", diagrams["roles"]),
    ]:
        story.append(Paragraph(title, styles["SectionTitle"]))
        story.append(_pdf_image(path, max_width=18 * cm, max_height=10 * cm))
        story.append(Spacer(1, 0.2 * cm))

    story.append(PageBreak())
    story.append(Paragraph("3. Data Quality, Repository, Finance, And AI Readiness", styles["SectionTitle"]))
    reference = stats["reference"]
    finance = stats["finance"]
    analytics = stats["nursing_analytics"]
    analytics_kpis = analytics.get("kpis", {})
    cases = stats["case_management"]
    engagement = stats["engagement"]
    story.append(_pdf_table(["Area", "Current position"], [
        ("Nursing analytics snapshot", f"{analytics_kpis.get('total_lifecycle_records', analytics.get('lifecycle_facts', 0))} lifecycle records; ATP {analytics_kpis.get('clean_atp_records', 0)}, provisional {analytics_kpis.get('clean_provisional_records', 0)}, full licence {analytics_kpis.get('clean_full_licence_records', 0)}"),
        ("Recognised PNG nursing schools", reference["png_nursing_school_count"]),
        ("Government / Non-government schools", f"{reference['government_nursing_school_count']} government, {reference['non_government_nursing_school_count']} non-government"),
        ("Open missing-data review items", MissingDataReview.objects.exclude(status="resolved").count()),
        ("Open duplicate review items", DuplicateReviewQueue.objects.filter(status="pending").count()),
        ("Repository folders / documents / versions / approvals", f"{stats['documents']['folders']} / {stats['documents']['documents']} / {stats['documents']['versions']} / {stats['documents']['approvals']}"),
        ("ICMS / Discipline / Decision records", f"{cases['complaints']} / {cases['disciplinary_cases']} / {cases['decisions']}"),
        ("Mapped entities / geocoded", f"{engagement['mapped_entities']} / {engagement['geocoded_entities']}"),
        ("Nursing Council finance combined total", f"PGK {finance.get('combined_total', 0)}"),
        ("AI Assistant mode", f"{stats['ai']['label']} - {stats['ai']['detail']}"),
    ], [6.1 * cm, 11.3 * cm]))

    story.append(PageBreak())
    story.append(Paragraph("4. Current System Interfaces", styles["SectionTitle"]))
    for item in screenshots:
        story.append(Paragraph(item["title"], styles["SectionTitle"]))
        story.append(Paragraph(item["description"], styles["BodyClean"]))
        story.append(_pdf_image(item["path"], max_width=18 * cm, max_height=10.4 * cm))
        story.append(PageBreak())

    doc.build(story, onFirstPage=_pdf_footer, onLaterPages=_pdf_footer)


def _docx_add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        cell.text = str(header)
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
        cell._tc.get_or_add_tcPr()
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cells[index].text = str(value)
    return table


def build_presentation_docx(stats, diagrams, screenshots):
    doc = WordDocument()
    section = doc.sections[0]
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)
    style = doc.styles["Normal"]
    style.font.name = "Aptos"
    style.font.size = Pt(10.5)

    if LOGO_PATH.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(LOGO_PATH), width=Inches(1.35))
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("PAPUA NEW GUINEA NATIONAL DEPARTMENT OF HEALTH\n")
    run.bold = True
    run.font.size = Pt(13)
    run = title.add_run(PROJECT_TITLE)
    run.bold = True
    run.font.size = Pt(14)
    doc.add_paragraph(f"Generated: {stats['generated_on']}")
    doc.add_paragraph("Subject: Presentation-ready briefing pack for management, staff training, ICT review, and launch readiness.")

    doc.add_heading("1. Executive Overview In Plain Language", level=1)
    doc.add_paragraph("The platform is a government regulatory operations system for the Nursing Council and Medical Board. It supports applications, registry records, licences, qualifications, documents, receipts, imports, data-quality review, dashboards, reports, and staff workflows.")
    doc.add_paragraph("Core rule: imported rows are not automatically trusted. They are staged, validated, cleansed, reviewed, approved, then promoted into live registry records.")
    doc.add_heading("Nursing Council cleansed analytics totals", level=2)
    doc.add_paragraph("These are the current Nursing Council figures after the cleanse. They come from the active cleansed analytics snapshot and should be used for the presentation executive overview. Legal person-table records remain protected until controlled workflow promotion.")
    _docx_add_table(doc, ["Statistic", "Current total", "Meaning"], [(label, _format_number(value), meaning) for label, value, meaning in stats["headline"]])
    doc.add_heading("Cleansed Nursing Council cadre / stage breakdown", level=2)
    _docx_add_table(
        doc,
        ["Cadre", "Provisional", "Full licence", "ATP", "Total"],
        [
            (
                row["cadre"],
                _format_number(row["provisional_licence_count"]),
                _format_number(row["full_licence_count"]),
                _format_number(row["authority_to_practice_count"]),
                _format_number(row["grand_total"]),
            )
            for row in stats["nursing_cadre_rows"][:10]
        ],
    )
    doc.add_heading("Separate operational live registry status", level=2)
    doc.add_paragraph("The following table is not the cleansed Nursing Council workforce total. It shows live legal person-table and platform operational counts while cleansed records are promoted through controlled workflow.")
    _docx_add_table(doc, ["Statistic", "Current total", "Meaning"], [(label, _format_number(value), meaning) for label, value, meaning in stats["operational_status"]])

    doc.add_heading("2. Department-Wide Diagrams", level=1)
    for title, path in [
        ("Enterprise Architecture", diagrams["architecture"]),
        ("Regulatory Workflow", diagrams["workflow"]),
        ("Data Governance Flow", diagrams["data_flow"]),
        ("Role Access And Privacy", diagrams["roles"]),
    ]:
        doc.add_heading(title, level=2)
        doc.add_picture(str(path), width=Inches(6.9))

    doc.add_heading("3. Current Interface Screenshots", level=1)
    for item in screenshots:
        doc.add_heading(item["title"], level=2)
        doc.add_paragraph(item["description"])
        doc.add_picture(str(item["path"]), width=Inches(6.9))

    doc.add_heading("4. Staff Presentation Notes", level=1)
    for item in [
        "Nursing Council and Medical Board workspaces must remain separated.",
        "Staff must clear missing-data and duplicate-review items before publishing management statistics.",
        "Documents should be uploaded into the correct office-scoped repository with metadata, versions, approval/rejection sign-off, and audit records.",
        "ICMS complaints, disciplinary cases, and regulatory decisions should be used for formal case-management and defensible outcomes.",
        "Nursing Council analytics snapshot rows are dashboard evidence, not legal registry identities.",
        "NHWA workbooks are reporting/sign-off outputs and must not overwrite registry records automatically.",
        "Mapped entities need verified stored coordinates before public demonstration.",
        "Financial reports must keep manual receipts and spreadsheet receipt rows visible as separate streams.",
        "The free local GPT assistant is guidance only and must not replace registrar approval or official workflow controls.",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    try:
        doc.save(PACK_DOCX)
    except PermissionError:
        fallback_docx = OUTPUT_DIR / f"NDOH_Regulatory_Platform_Presentation_Brief_{DATE_STAMP}_updated_{datetime.now():%Y%m%d%H%M%S}.docx"
        doc.save(fallback_docx)
        print(f"Saved updated Word brief to {fallback_docx} because {PACK_DOCX.name} is open or locked.")


def build_index_pdf(stats):
    doc = SimpleDocTemplate(str(INDEX_PDF), pagesize=A4, rightMargin=1.35 * cm, leftMargin=1.35 * cm, topMargin=1.35 * cm, bottomMargin=1.6 * cm)
    rows = [
        ("Presentation pack PDF", PACK_PDF.name, "Main official presentation document with diagrams, screenshots, and live statistics."),
        ("Presentation brief Word", PACK_DOCX.name, "Editable Word version for management comments and submission formatting."),
        ("Presentation pack Markdown", PACK_MD.name, "Plain text source for quick edits and audit trail."),
        ("Screenshot folder", "assets/screenshots", "Fresh interface screenshots generated for this pack."),
        ("Diagram folder", "assets/diagrams", "Enterprise architecture, workflow, data governance, and privacy diagrams."),
        ("Full-scope user guide", "docs/NDOH_Full_Scope_Platform_User_Guide_20260601.pdf", "Role-based staff guide."),
        ("Current platform update brief", "docs/PLATFORM_UPDATE_BRIEF_20260601.md", "Latest interface and workflow update summary."),
        ("Documentation index", "DOCUMENTATION_INDEX.md", "Master documentation register."),
        ("Government launch package", "docs/government_launch_package", "Architecture, security, data governance, deployment, testing, and AI setup documents."),
    ]
    story = []
    if LOGO_PATH.exists():
        story.append(PdfImage(str(LOGO_PATH), width=2.4 * cm, height=2.0 * cm))
    story.append(Paragraph("Presentation Documentation Index", styles["CoverTitle"]))
    story.append(Paragraph(PROJECT_TITLE, styles["BodyClean"]))
    story.append(Paragraph(f"Generated: {stats['generated_on']}", styles["BodyClean"]))
    story.append(Spacer(1, 0.4 * cm))
    story.append(_pdf_table(["Document", "Location", "Purpose"], rows, [4.5 * cm, 5.8 * cm, 7.1 * cm]))
    doc.build(story, onFirstPage=_pdf_footer, onLaterPages=_pdf_footer)


def write_readme():
    README.write_text(
        f"""# NDOH Regulatory Platform Presentation Pack

Generated: {DISPLAY_DATE}

This folder contains the official presentation-ready documentation pack.

- `{PACK_PDF.name}` - PDF presentation pack with live statistics, diagrams, workflows, and screenshots.
- `{PACK_DOCX.name}` - editable Word briefing version.
- `{PACK_MD.name}` - Markdown source.
- `{INDEX_PDF.name}` - documentation index PDF.
- `WHO_Conference_Room_Meeting_Prep_20260604.md` - focused two-hour meeting run sheet, demo script, talking points, expected questions, and pre-meeting checklist for the Thursday 4 June 2026 WHO Conference Room presentation.
- `../PLATFORM_UPDATE_BRIEF_20260601.md` - current platform update brief for the latest analytics snapshot, ICMS, discipline, decisions, NHWA, mapping, receipt linking, UI, notification, Records Hub, duplicate review, and registration changes.
- `assets/screenshots/` - refreshed interface screenshots.
- `assets/diagrams/` - enterprise architecture and workflow diagrams.

The screenshots were generated from current system views and current database statistics using the local project environment.
""",
        encoding="utf-8",
    )


def main():
    _ensure_dirs()
    diagrams = create_diagrams()
    screenshots = capture_screenshots()
    stats = collect_live_statistics()
    build_markdown(stats, diagrams, screenshots)
    build_presentation_pdf(stats, diagrams, screenshots)
    build_presentation_docx(stats, diagrams, screenshots)
    build_index_pdf(stats)
    write_readme()
    print(PACK_PDF)
    print(PACK_DOCX)
    print(INDEX_PDF)
    print(PACK_MD)


if __name__ == "__main__":
    main()
