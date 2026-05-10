from pathlib import Path
import csv
import re

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


PROJECT_NAME = "The National Department Of Health Regulatory Bodies Nursing Council & The Medical Board Online Workforce System"
SUBJECT = f"User Management Brief for {PROJECT_NAME}."
BASE_DIR = Path(__file__).resolve().parent.parent
DOCS_DIR = BASE_DIR / "docs"
NOTEBOOKS_DIR = BASE_DIR / "notebooks"
ASSETS_DIR = DOCS_DIR / "system_brief_assets"
OUTPUT_PATH = DOCS_DIR / "NDOH_Regulatory_Bodies_Online_Workforce_System_Brief.docx"


def read_summary(path: Path):
    result = {}
    if not path.exists():
        return result
    for line in path.read_text(encoding="utf-8").splitlines():
        if ":" not in line:
            continue
        key, value = line.split(":", 1)
        result[key.strip()] = value.strip()
    return result


def read_issue_counts(path: Path, issue_column: str):
    counts = {}
    if not path.exists():
        return counts
    with path.open(newline="", encoding="utf-8") as handle:
        reader = csv.DictReader(handle)
        for row in reader:
            issue = (row.get(issue_column) or "").strip()
            if not issue:
                continue
            counts[issue] = counts.get(issue, 0) + 1
    return counts


def set_default_font(document: Document):
    style = document.styles["Normal"]
    style.font.name = "Times New Roman"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    style.font.size = Pt(11)

    for style_name in ["Title", "Heading 1", "Heading 2", "Heading 3"]:
        if style_name in document.styles:
            style = document.styles[style_name]
            style.font.name = "Times New Roman"
            style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")


def add_bottom_border(paragraph):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "18")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "000000")
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def set_cell_width(cell, width_inches):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(width_inches * 1440)))
    tc_w.set(qn("w:type"), "dxa")


def add_bullet(document: Document, text: str):
    paragraph = document.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.add_run(text)


def add_number(document: Document, text: str):
    paragraph = document.add_paragraph(style="List Number")
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.add_run(text)


def add_interface_row(document: Document, title: str, image_a: str, caption_a: str, image_b: str, caption_b: str, explanation_lines):
    document.add_paragraph()
    heading = document.add_paragraph()
    heading_run = heading.add_run(title)
    heading_run.bold = True
    heading_run.font.size = Pt(12)

    table = document.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    widths = [2.15, 2.15, 2.9]
    for cell, width in zip(table.rows[0].cells, widths):
        set_cell_width(cell, width)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP

    left = table.rows[0].cells[0]
    p = left.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(ASSETS_DIR / image_a), width=Inches(2.0))
    caption = left.add_paragraph(caption_a)
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.runs[0].italic = True
    caption.runs[0].font.size = Pt(9)

    middle = table.rows[0].cells[1]
    p = middle.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(ASSETS_DIR / image_b), width=Inches(2.0))
    caption = middle.add_paragraph(caption_b)
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.runs[0].italic = True
    caption.runs[0].font.size = Pt(9)

    right = table.rows[0].cells[2]
    right.paragraphs[0].add_run("What this interface shows").bold = True
    for line in explanation_lines:
        bullet = right.add_paragraph(style="List Bullet")
        bullet.paragraph_format.space_after = Pt(0)
        bullet.add_run(line)


def build_document():
    doc = Document()
    set_default_font(doc)

    section = doc.sections[0]
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)

    crest = DOCS_DIR.parent / "static" / "img" / "NDOH_LOGO.png"
    if crest.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(crest), width=Inches(1.8))

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("PAPUA NEW GUINEA NURSING COUNCIL\nOFFICE OF THE REGISTRAR")
    run.bold = True
    run.font.size = Pt(15)

    border = doc.add_paragraph()
    border.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_bottom_border(border)

    doc.add_paragraph()
    salutation = doc.add_paragraph()
    salutation.add_run("Dear Ms. Romanah,").bold = True

    subject = doc.add_paragraph()
    subject_run = subject.add_run(f"Subject: {SUBJECT}")
    subject_run.bold = True
    subject_run.underline = True

    intro = [
        f"This brief provides a plain-language overview of {PROJECT_NAME}. It is written for managers, operational staff, and decision-makers who need to understand what the system does, who can access it, how privacy is controlled, and what is required for reliable data capture from paper records to the live electronic register.",
        "The document covers the full scope of the project: public entry screens, staff dashboards, Nursing Council functions, Medical Board functions, self-service user portals, reporting features, record management, privacy controls, and data-cleansing observations.",
    ]
    for text in intro:
        paragraph = doc.add_paragraph(text)
        paragraph.paragraph_format.space_after = Pt(6)

    doc.add_heading("1. System Name and Purpose", level=1)
    doc.add_paragraph(
        f"The official project name used in this brief is: {PROJECT_NAME}. "
        "In simple terms, the system is an online workforce management and registration platform for the Nursing Council and the Medical Board under the National Department of Health."
    )
    doc.add_paragraph(
        "It is designed to help the offices receive applications, register health professionals, track licence history, store qualifications and supporting documents, review missing information, and generate management dashboards and reports."
    )

    doc.add_heading("2. Executive Summary in Plain Language", level=1)
    for item in [
        "The system is both a day-to-day working tool and a long-term record system.",
        "It separates Nursing Council and Medical Board information so each office only sees the data it is allowed to view.",
        "It supports staff work, public form intake, and role-based user portals for professionals and graduands.",
        "It can store profile details, licences, qualifications, documents, payments, applications, historical imports, and data-quality issues.",
        "The biggest challenge is not the screens themselves. The biggest challenge is the quality of raw source records coming from paper files and old spreadsheets.",
    ]:
        add_bullet(doc, item)

    doc.add_heading("3. Main Interfaces With Screenshots", level=1)
    add_interface_row(
        doc,
        "3.1 Public entry and leadership overview",
        "public_home.png",
        "Public home page",
        "overall_dashboard_admin.png",
        "Overall dashboard",
        [
            "The public home page is the front door of the system. It helps users find the right pathway into the registry.",
            "The overall dashboard gives a combined high-level view for leadership and approved staff roles.",
            "This interface supports a big-picture view of registrations, applications, and workforce information across the platform.",
        ],
    )
    add_interface_row(
        doc,
        "3.2 Separate regulatory body dashboards",
        "nursing_dashboard.png",
        "Nursing Council dashboard",
        "medical_dashboard.png",
        "Medical Board dashboard",
        [
            "These are the main operational dashboards for the two regulatory bodies.",
            "Nursing Council staff are expected to see nursing, midwifery, nurse aide, and graduand information only.",
            "Medical Board staff are expected to see medical doctor, CHW, and medical-board form information only.",
            "The system now enforces this separation across dashboards, records, application review, exports, and related API access.",
        ],
    )
    add_interface_row(
        doc,
        "3.3 Application intake and form guidance",
        "nursing_forms.png",
        "Nursing Council forms portal",
        "medical_forms.png",
        "Medical Board forms portal",
        [
            "These screens guide applicants or staff to the correct official form pathway.",
            "They reduce errors by separating nursing forms from medical-board forms.",
            "They allow the system to capture the right documents, details, and workflow for each application type.",
        ],
    )
    add_interface_row(
        doc,
        "3.4 Self-service professional portals",
        "nurse_portal.png",
        "Nurse portal",
        "doctor_portal.png",
        "Doctor portal",
        [
            "These are individual user dashboards for professionals rather than staff registrars.",
            "They are intended for personal record viewing, application tracking, receipt submission, and limited self-service access.",
            "A personal user should only see their own information, not another person’s record.",
        ],
    )
    add_interface_row(
        doc,
        "3.5 Graduand and individual record management",
        "graduand_portal.png",
        "Graduand portal",
        "professional_detail.png",
        "Professional record detail",
        [
            "The graduand portal supports students and future applicants as they move into professional registration.",
            "The professional detail screen is the core individual file, bringing together profile details, qualifications, documents, images, and related records.",
            "This is important because staff can review one person’s complete file in one place.",
        ],
    )
    add_interface_row(
        doc,
        "3.6 Workflow monitoring and administration",
        "workforce_flow.png",
        "Workforce flow dashboard",
        "admin_dashboard.png",
        "Admin dashboard",
        [
            "The workforce flow view supports planning and trend analysis by year, pathway, and distribution.",
            "The admin dashboard supports senior staff or administrators who need broad operational oversight.",
            "These interfaces support governance, reporting, monitoring, and follow-up work.",
        ],
    )

    doc.add_heading("4. Roles, Accessibilities, and Privacy", level=1)
    doc.add_paragraph(
        "The system is role-based. This means every user does not see the same information. Access depends on the person’s approved role and whether they belong to Nursing Council, Medical Board, or a personal self-service pathway."
    )
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["Role", "Main access", "What the role can view", "Privacy position"]
    for cell, text in zip(table.rows[0].cells, headers):
        cell.text = text

    role_rows = [
        ("Admin", "Full system oversight", "Can access both Nursing Council and Medical Board functions, reports, dashboards, and records.", "Cross-domain access is allowed because this is the top supervisory role."),
        ("Nursing Council registrar / staff", "Nursing Council staff dashboards", "Can view nursing dashboards, nursing applications, nursing forms, nursing records, nursing reports, and nursing data quality work.", "Should not view Medical Board doctor or CHW data."),
        ("Medical Board registrar / staff", "Medical Board staff dashboards", "Can view medical dashboards, medical forms, medical applications, doctor and CHW records, and medical reporting.", "Should not view Nursing Council data."),
        ("Nurse / Nurse Aide / Doctor / CHW", "Personal portal", "Can view their own records, applications, and personal tasks relevant to their role.", "Should not view another professional’s record."),
        ("Graduand", "Student / graduand portal", "Can view graduand information, application status, receipts, and related pathway actions.", "Should only see their own personal pipeline data."),
        ("Viewer / Reviewer", "Limited support access", "Can be assigned controlled access for review or viewing purposes depending on configuration.", "Should stay within assigned workflow only."),
    ]
    for row in role_rows:
        cells = table.add_row().cells
        for cell, value in zip(cells, row):
            cell.text = value

    doc.add_paragraph()
    privacy = doc.add_paragraph()
    privacy.add_run("Important privacy note: ").bold = True
    privacy.add_run(
        "The system now separates Nursing Council and Medical Board data at dashboard level, record level, application level, report/export level, and staff API level. This is a key requirement for confidentiality and operational discipline."
    )

    doc.add_heading("5. What the System Can Do", level=1)
    for item in [
        "Receive online applications through separated Nursing Council and Medical Board form pathways.",
        "Maintain workforce profiles for nurses, midwives, nurse aides, doctors, CHWs, and graduands.",
        "Store qualification history, institution details, completion year, and country of study.",
        "Store images and uploaded supporting documents such as IDs, certificates, and transcripts.",
        "Track applications by form code and status such as pending, approved, or rejected.",
        "Record receipt submissions and payment evidence for applications.",
        "Show dashboards for staffing, workflow, yearly trends, province distribution, and licence movement.",
        "Import historical workbook records and convert them into structured electronic records.",
        "Flag incomplete or suspicious records for data quality review and follow-up.",
        "Support management reporting through dashboards and export functions.",
    ]:
        add_bullet(doc, item)

    doc.add_heading("6. What the System Is Able to Capture", level=1)
    capture_table = doc.add_table(rows=1, cols=2)
    capture_table.style = "Table Grid"
    capture_table.rows[0].cells[0].text = "Data area"
    capture_table.rows[0].cells[1].text = "Examples of information captured"
    capture_rows = [
        ("Identity and contacts", "Names, registration number, email, phone, gender, date of birth, province."),
        ("Professional classification", "Nurse, midwife, nurse aide, doctor, CHW, graduand, applicant type, cadre."),
        ("Licence information", "Issue date, expiry date, licence pathway, provisional or full registration status."),
        ("Education and qualifications", "Institution, qualification title, programme completed, completion year, country."),
        ("Supporting records", "Certificates, transcripts, ID documents, uploaded files, receipt images, photos."),
        ("Applications", "Form code, pathway, status, submitted date, review notes."),
        ("Historical workforce records", "Imported workbook licence rows, yearly statistics, payment-linked records."),
        ("Data quality tracking", "Missing fields, severity level, review status, notifications."),
    ]
    for left, right in capture_rows:
        row = capture_table.add_row().cells
        row[0].text = left
        row[1].text = right

    doc.add_heading("7. Data Cleansing Findings and Current Challenges", level=1)
    full_summary = read_summary(NOTEBOOKS_DIR / "full_registrations_summary.txt")
    provisional_summary = read_summary(NOTEBOOKS_DIR / "provisional_graduands_summary.txt")
    full_issues = read_issue_counts(NOTEBOOKS_DIR / "full_registrations_quality_issues.csv", "issue_detail")
    prov_issues = read_issue_counts(NOTEBOOKS_DIR / "provisional_graduands_quality_issues.csv", "issues")

    doc.add_paragraph(
        "The project already includes data-cleansing and auditing logic. This means the system does not simply copy raw workbook or paper information into the database without checks. It tries to normalize names, numbers, dates, institutions, and provinces, and it flags records that need further review."
    )
    metrics = [
        f"Full-registration records cleaned: {full_summary.get('Cleaned records', 'N/A')}",
        f"Full-registration issue rows flagged: {full_summary.get('Quality issue rows', 'N/A')}",
        f"Provisional / graduand records cleaned: {provisional_summary.get('Cleaned records', 'N/A')}",
        f"Provisional / graduand issue rows flagged: {provisional_summary.get('Quality issue records', 'N/A')}",
    ]
    for metric in metrics:
        add_bullet(doc, metric)

    doc.add_paragraph("The main data problems currently being found include:")
    problem_list = [
        f"Invalid practitioner number formats ({full_issues.get('Invalid practitioner number format not imported', 0)} flagged rows in the full-registration issue file).",
        f"Duplicate practitioner numbers ({full_issues.get('Duplicate practitioner number not imported', 0)} flagged rows in the full-registration issue file).",
        f"Missing or invalid issued dates ({full_issues.get('Missing or invalid issued date', 0)} in the full-registration issue file, plus provisional date issues such as {prov_issues.get('Missing or invalid issued date; Issued date status: missing_issued_date', 0)} rows in the provisional issue file).",
        f"Missing graduation year ({full_issues.get('Missing graduation year', 0)} full-registration flags and several provisional flags).",
        f"Missing institution details ({full_issues.get('Missing institution', 0)} flagged rows in the full-registration issue file).",
        "Conflicting names sharing the same registration number.",
        "Old spreadsheet dates that needed repair because the year or text format was clearly wrong.",
    ]
    for item in problem_list:
        add_bullet(doc, item)

    doc.add_paragraph(
        "In simple language, this means some source records do not match each other, some are incomplete, and some are not reliable enough to enter directly into the final register without review."
    )

    doc.add_heading("8. What Must Be Done for Clean Data Flow From Paper to Electronic Form", level=1)
    actions = [
        "Use one standard paper intake checklist for every application package.",
        "Assign a unique intake reference number as soon as a paper file is received.",
        "Scan the complete paper file early, including form, ID, qualification documents, and receipt.",
        "Enter records first into a staging or review step rather than directly into the final live register.",
        "Make critical fields mandatory before approval: name, registration number, pathway, institution, issue date, and province.",
        "Use dropdown lists for provinces, institutions, cadres, and document types instead of free typing whenever possible.",
        "Run duplicate checks against registration number, practitioner number, name, and email before approval.",
        "Require registrar review against the source scan before moving a record into the official register.",
        "Use the missing-data review tools regularly so incomplete records are corrected quickly.",
        "Carry out routine monthly or quarterly quality audits instead of waiting for one large cleanup exercise.",
    ]
    for item in actions:
        add_number(doc, item)

    doc.add_heading("9. Overall Conclusion", level=1)
    conclusion = [
        f"{PROJECT_NAME} is already capable of managing much more than simple online registration.",
        "It supports separated regulatory body operations, role-based privacy, professional self-service views, qualification and document storage, historical import handling, workflow monitoring, and management reporting.",
        "The strongest remaining need is to improve raw data quality from paper files and historical sources so the electronic register remains accurate, trusted, and sustainable over time.",
    ]
    for text in conclusion:
        doc.add_paragraph(text)

    doc.save(OUTPUT_PATH)
    return OUTPUT_PATH


if __name__ == "__main__":
    output = build_document()
    print(output)
